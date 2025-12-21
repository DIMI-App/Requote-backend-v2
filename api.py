from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import json
import subprocess
from werkzeug.utils import secure_filename
import threading
import time
import shutil

# Import the Python converter
from python_converter_final import convert_to_pdf_python

app = Flask(__name__)

CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization", "Accept"],
        "expose_headers": ["Content-Type"],
        "supports_credentials": False,
        "max_age": 3600
    }
})

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# Allowed file extensions
ALLOWED_OFFER1_EXTENSIONS = {'pdf', 'docx', 'doc', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'}
ALLOWED_OFFER2_EXTENSIONS = {'docx', 'doc', 'xlsx', 'xls', 'pdf'}

# Thread-safe status storage
_status_lock = threading.Lock()
processing_status = {
    'status': 'idle',
    'message': '',
    'items_count': 0,
    'items': [],
    'started_at': None,
    'updated_at': None,
    'file_format': None
}

def allowed_file(filename, allowed_extensions):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

def get_file_extension(filename):
    """Get file extension in lowercase"""
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def convert_to_docx_python(input_path, output_path, file_format):
    """Convert template formats to DOCX using Python"""
    try:
        print(f"Converting {file_format.upper()} template to DOCX...", flush=True)
        
        if file_format == 'docx':
            shutil.copy(input_path, output_path)
            print("✓ DOCX template ready", flush=True)
            return True
        
        elif file_format == 'pdf':
            import fitz
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            print("Converting PDF template to DOCX with table extraction...", flush=True)
            pdf_doc = fitz.open(input_path)
            docx_doc = Document()
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                tables = page.find_tables()
                
                if tables:
                    print(f"  Found {len(tables)} table(s) on page {page_num + 1}", flush=True)
                    for table in tables:
                        table_data = table.extract()
                        if not table_data or len(table_data) == 0:
                            continue
                        
                        num_rows = len(table_data)
                        num_cols = max(len(row) for row in table_data) if table_data else 0
                        
                        if num_cols > 0:
                            docx_table = docx_doc.add_table(rows=num_rows, cols=num_cols)
                            docx_table.style = 'Light Grid Accent 1'
                            
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_text in enumerate(row_data):
                                    if col_idx < num_cols:
                                        cell = docx_table.rows[row_idx].cells[col_idx]
                                        cell.text = str(cell_text) if cell_text else ""
                else:
                    text = page.get_text()
                    if text.strip():
                        para = docx_doc.add_paragraph(text)
                
                if page_num < len(pdf_doc) - 1:
                    docx_doc.add_page_break()
            
            pdf_doc.close()
            docx_doc.save(output_path)
            print("✓ PDF converted to DOCX template with tables", flush=True)
            return True
        
        elif file_format == 'doc':
            print("✗ DOC format requires LibreOffice. Please upload DOCX.", flush=True)
            return False
            
        elif file_format in ['xlsx', 'xls']:
            import openpyxl
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print(f"Converting {file_format.upper()} template to DOCX...", flush=True)
            workbook = openpyxl.load_workbook(input_path, data_only=True)
            sheet = workbook.active
            docx_doc = Document()
            
            all_rows = []
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                all_rows.append(row_data)
            
            print(f"  Found {len(all_rows)} rows in Excel", flush=True)
            
            table_start_row = None
            for idx, row in enumerate(all_rows):
                row_text = ' '.join(row).upper()
                if any(keyword in row_text for keyword in ['POSITION', 'DESCRIPTION', 'PRICE', 'QUANTITY', 'TOTAL']):
                    table_start_row = idx
                    break
            
            if table_start_row is None:
                table_start_row = 0
            
            for idx in range(table_start_row):
                row_text = ' '.join(all_rows[idx]).strip()
                if row_text:
                    para = docx_doc.add_paragraph(row_text)
            
            if table_start_row > 0:
                docx_doc.add_paragraph()
            
            table_rows = all_rows[table_start_row:]
            
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                num_rows = len(table_rows)
                
                docx_table = docx_doc.add_table(rows=num_rows, cols=num_cols)
                docx_table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx in range(num_cols):
                        cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                        docx_table.rows[row_idx].cells[col_idx].text = cell_text
                        
                        if row_idx == 0:
                            for paragraph in docx_table.rows[row_idx].cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            docx_doc.save(output_path)
            print(f"✓ {file_format.upper()} converted to DOCX template", flush=True)
            return True
            
        else:
            print(f"✗ Unsupported template format: {file_format}", flush=True)
            return False
            
    except Exception as e:
        print(f"✗ Template conversion error: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

@app.route('/', methods=['GET'])
def home():
    return jsonify({
        'message': 'Requote AI Backend is running!',
        'version': 'SV15-Offer3-Generation',
        'status': 'healthy',
        'supported_formats': {
            'offer1': list(ALLOWED_OFFER1_EXTENSIONS),
            'offer2': list(ALLOWED_OFFER2_EXTENSIONS)
        }
    })

def process_file_background(filepath, file_extension):
    """Background processing using semantic extraction"""
    global processing_status
    
    try:
        with _status_lock:
            processing_status['status'] = 'processing'
            processing_status['message'] = f'Processing {file_extension.upper()} file...'
            processing_status['file_format'] = file_extension
            processing_status['started_at'] = time.time()
            processing_status['updated_at'] = time.time()
        
        print("=== BACKGROUND PROCESSING STARTED ===", flush=True)
        
        pdf_path = os.path.join(UPLOAD_FOLDER, 'offer1.pdf')
        
        if file_extension == 'pdf':
            if filepath != pdf_path:
                shutil.copy(filepath, pdf_path)
        else:
            with _status_lock:
                processing_status['message'] = f'Converting {file_extension.upper()} to PDF...'
            
            conversion_success = convert_to_pdf_python(filepath, pdf_path, file_extension)
            if not conversion_success:
                with _status_lock:
                    processing_status['status'] = 'error'
                    processing_status['message'] = f'Failed to convert {file_extension.upper()}'
                return
        
        with _status_lock:
            processing_status['message'] = 'Extracting items with semantic analysis...'
        
        items_output_path = os.path.join(OUTPUT_FOLDER, 'items_offer1.json')
        extract_script_path = os.path.join(BASE_DIR, 'extract_pdf_direct_enhanced.py')
        
        result = subprocess.run(
            ['python', extract_script_path],
            capture_output=True,
            text=True,
            cwd=BASE_DIR,
            timeout=300
        )
        
        if result.returncode != 0 or not os.path.exists(items_output_path):
            with _status_lock:
                processing_status['status'] = 'error'
                processing_status['message'] = 'Extraction failed'
            return
        
        with open(items_output_path, 'r', encoding='utf-8') as f:
            full_data = json.load(f)
        
        items = full_data.get('items', [])
        
        with _status_lock:
            processing_status['status'] = 'completed'
            processing_status['message'] = f'Successfully extracted {len(items)} items'
            processing_status['items_count'] = len(items)
            processing_status['items'] = items
            processing_status['updated_at'] = time.time()
        
    except Exception as e:
        print(f"=== ERROR: {str(e)} ===", flush=True)
        with _status_lock:
            processing_status['status'] = 'error'
            processing_status['message'] = f'Error: {str(e)}'

@app.route('/api/process-offer1', methods=['POST', 'OPTIONS'])
def api_process_offer1():
    """Process Offer 1 - semantic extraction"""
    global processing_status
    
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        print("=" * 60, flush=True)
        print("Received request to process Offer 1", flush=True)
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename, ALLOWED_OFFER1_EXTENSIONS):
            return jsonify({
                'error': f'Unsupported file format. Allowed: {", ".join(ALLOWED_OFFER1_EXTENSIONS)}'
            }), 400
        
        file_extension = get_file_extension(file.filename)
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, f'offer1_original.{file_extension}')
        file.save(filepath)
        
        print(f"✓ File saved: {filepath}", flush=True)
        print(f"✓ Format: {file_extension.upper()}", flush=True)
        
        # Reset status
        with _status_lock:
            processing_status = {
                'status': 'processing',
                'message': f'File uploaded, starting processing...',
                'items_count': 0,
                'items': [],
                'file_format': file_extension,
                'started_at': time.time(),
                'updated_at': time.time()
            }
        
        # Start background thread
        thread = threading.Thread(
            target=process_file_background,
            args=(filepath, file_extension),
            daemon=True
        )
        
        thread.start()
        
        print(f"✓ Background thread started", flush=True)
        print("=" * 60, flush=True)
        
        return jsonify({
            'success': True,
            'message': 'Processing started. Poll /api/status for updates.',
            'status': 'processing',
            'file_format': file_extension
        })
        
    except Exception as e:
        print(f"ERROR in process-offer1: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/status', methods=['GET', 'OPTIONS'])
def api_status():
    if request.method == 'OPTIONS':
        return '', 204
    
    with _status_lock:
        status_copy = processing_status.copy()
    
    if status_copy.get('started_at') and status_copy['status'] == 'processing':
        elapsed = time.time() - status_copy['started_at']
        status_copy['elapsed_seconds'] = round(elapsed, 1)
    
    return jsonify(status_copy)

@app.route('/api/upload-offer2', methods=['POST', 'OPTIONS'])
def api_upload_offer2():
    """Upload Offer 2 template and extract company branding"""
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        print("=" * 60, flush=True)
        print("Received Offer 2 template", flush=True)
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename, ALLOWED_OFFER2_EXTENSIONS):
            return jsonify({
                'error': f'Unsupported template format. Allowed: {", ".join(ALLOWED_OFFER2_EXTENSIONS)}'
            }), 400
        
        file_extension = get_file_extension(file.filename)
        
        print(f"✓ Template format: {file_extension.upper()}", flush=True)
        
        # Clean up old template files
        old_docx = os.path.join(BASE_DIR, 'offer2_template.docx')
        old_xlsx = os.path.join(BASE_DIR, 'offer2_template.xlsx')
        old_xls = os.path.join(BASE_DIR, 'offer2_template.xls')
        
        for old_file in [old_docx, old_xlsx, old_xls]:
            if os.path.exists(old_file):
                os.remove(old_file)
        
        # Save template as DOCX (always convert to DOCX for company extraction)
        template_path = os.path.join(BASE_DIR, 'offer2_template.docx')
        
        if file_extension == 'docx':
            file.save(template_path)
        else:
            # Convert to DOCX first
            original_path = os.path.join(BASE_DIR, f'offer2_template_original.{file_extension}')
            file.save(original_path)
            
            conversion_success = convert_to_docx_python(original_path, template_path, file_extension)
            
            if not conversion_success:
                return jsonify({
                    'error': f'Cannot convert {file_extension.upper()} to extract company data.',
                    'suggestion': 'Please try uploading a DOCX template instead.'
                }), 500
        
        print(f"✓ Template saved: {template_path}", flush=True)
        
        # Extract company data from template
        print("Extracting company data from template...", flush=True)
        
        extract_script_path = os.path.join(BASE_DIR, 'extract_company_data.py')
        
        result = subprocess.run(
            ['python', extract_script_path],
            capture_output=True,
            text=True,
            cwd=BASE_DIR,
            timeout=180
        )
        
        if result.stdout:
            print(result.stdout, flush=True)
        if result.stderr:
            print(result.stderr, flush=True)
        
        company_data_path = os.path.join(OUTPUT_FOLDER, 'company_data.json')
        
        if result.returncode != 0 or not os.path.exists(company_data_path):
            print("⚠ Company extraction had issues, but continuing...", flush=True)
            # Don't fail - we can still generate basic offer
        else:
            print("✓ Company data extracted successfully", flush=True)
        
        print("=" * 60, flush=True)
        
        return jsonify({
            'success': True,
            'message': f'Template uploaded and processed ({file_extension.upper()})',
            'file_format': file_extension,
            'company_extracted': os.path.exists(company_data_path)
        })
        
    except Exception as e:
        print(f"Error in upload-offer2: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-offer', methods=['POST', 'OPTIONS'])
def api_generate_offer():
    """Generate Offer 3 - NEW APPROACH: Build from scratch"""
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        print("=" * 60, flush=True)
        print("Starting Offer 3 generation (NEW APPROACH)", flush=True)
        print("=" * 60, flush=True)
        
        data = request.get_json() or {}
        markup = data.get('markup', 0)
        
        items_path = os.path.join(OUTPUT_FOLDER, 'items_offer1.json')
        if not os.path.exists(items_path):
            return jsonify({'error': 'No items found. Please process Offer 1 first.'}), 400
        
        # Apply markup if needed
        if markup > 0:
            print(f"Applying {markup}% markup...", flush=True)
            with open(items_path, 'r', encoding='utf-8') as f:
                full_data = json.load(f)
            
            items = full_data.get('items', [])
            items = apply_markup_to_items(items, markup)
            full_data['items'] = items
            
            with open(items_path, 'w', encoding='utf-8') as f:
                json.dump(full_data, f, ensure_ascii=False, indent=2)
        
        # Clean up old output files
        old_offer3 = os.path.join(OUTPUT_FOLDER, 'final_offer3.docx')
        old_offer1 = os.path.join(OUTPUT_FOLDER, 'final_offer1.docx')
        old_xlsx = os.path.join(OUTPUT_FOLDER, 'final_offer1.xlsx')
        
        for old_file in [old_offer3, old_offer1, old_xlsx]:
            if os.path.exists(old_file):
                os.remove(old_file)
        
        # Generate using NEW build_offer3.py script
        print("Building Offer 3 from scratch...", flush=True)
        build_script_path = os.path.join(BASE_DIR, 'build_offer3.py')
        
        result = subprocess.run(
            ['python', build_script_path],
            capture_output=True,
            text=True,
            cwd=BASE_DIR,
            timeout=120
        )
        
        if result.stdout:
            print(result.stdout, flush=True)
        if result.stderr:
            print(result.stderr, flush=True)
        
        if result.returncode != 0:
            return jsonify({
                'error': 'Offer generation failed',
                'details': result.stderr
            }), 500
        
        output_path = os.path.join(OUTPUT_FOLDER, 'final_offer3.docx')
        
        if not os.path.exists(output_path):
            return jsonify({'error': 'Output file not generated'}), 500
        
        with open(items_path, 'r', encoding='utf-8') as f:
            full_data = json.load(f)
            items = full_data.get('items', [])
        
        print(f"✓ Offer 3 generated successfully", flush=True)
        print("=" * 60, flush=True)
        
        return jsonify({
            'success': True,
            'message': 'Offer 3 generated successfully',
            'download_url': '/api/download-offer',
            'items_count': len(items),
            'output_format': 'docx',
            'generation_method': 'build_from_scratch'
        })
        
    except Exception as e:
        print(f"Error in generate-offer: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-offer', methods=['GET', 'OPTIONS'])
def api_download_offer():
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        # Check for new Offer 3 output first
        output_offer3 = os.path.join(OUTPUT_FOLDER, 'final_offer3.docx')
        output_offer1 = os.path.join(OUTPUT_FOLDER, 'final_offer1.docx')
        output_xlsx = os.path.join(OUTPUT_FOLDER, 'final_offer1.xlsx')
        
        if os.path.exists(output_offer3):
            output_path = output_offer3
            download_name = 'requoted_offer.docx'
        elif os.path.exists(output_xlsx):
            output_path = output_xlsx
            download_name = 'requoted_offer.xlsx'
        elif os.path.exists(output_offer1):
            output_path = output_offer1
            download_name = 'requoted_offer.docx'
        else:
            return jsonify({'error': 'No offer generated yet'}), 404
        
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        if output_path.endswith('.xlsx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        from flask import Response
        
        with open(output_path, 'rb') as f:
            file_data = f.read()
        
        response = Response(
            file_data,
            mimetype=mimetype,
            headers={
                'Content-Disposition': f'attachment; filename="{download_name}"',
                'Content-Type': mimetype,
                'Cache-Control': 'no-cache',
                'Pragma': 'no-cache',
                'Expires': '0'
            }
        )
        
        return response
        
    except Exception as e:
        print(f"Error in download-offer: {str(e)}", flush=True)
        return jsonify({'error': str(e)}), 500

def apply_markup_to_items(items, markup_percent):
    import re
    
    for item in items:
        price_str = str(item.get('unit_price', ''))
        if not price_str or price_str == '':
            price_str = str(item.get('price', ''))
        
        numbers = re.findall(r'\d+\.?\d*', price_str)
        if numbers:
            original_price = float(numbers[0])
            new_price = original_price * (1 + markup_percent / 100)
            currency = re.findall(r'[€$£¥]', price_str)
            currency_symbol = currency[0] if currency else '€'
            item['unit_price'] = currency_symbol + str(round(new_price, 2))
            if 'price' in item:
                item['price'] = currency_symbol + str(round(new_price, 2))
    
    return items

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    print("=" * 60)
    print("Starting Requote AI Backend - SV15 Offer 3 Generation")
    print(f"Server at: http://0.0.0.0:{port}")
    print(f"Supported formats: {', '.join(ALLOWED_OFFER1_EXTENSIONS)}")
    print("NEW: Builds Offer 3 from scratch instead of editing templates")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=port)
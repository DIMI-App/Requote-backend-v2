"""
Build Offer 3 - USE AI-EXTRACTED DESCRIPTIONS
1. Copy template (preserves header/footer/logo)
2. Clear body
3. Add pricing table
4. Add technical descriptions from JSON (AI-extracted)
5. Add commercial terms
"""

import os
import sys
import json
import shutil
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from standard_template import Offer3Template

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')

def generate_offer3(company_data_path, items_data_path, output_path):
    """
    Build Offer 3 using AI-extracted data
    """
    
    try:
        print("=" * 60, flush=True)
        print("BUILDING OFFER 3 - AI EXTRACTION APPROACH", flush=True)
        print("=" * 60, flush=True)
        
        # Load items data
        print(f"Loading items data: {items_data_path}", flush=True)
        if not os.path.exists(items_data_path):
            print("✗ Items data not found!", flush=True)
            return False
        
        with open(items_data_path, 'r', encoding='utf-8') as f:
            items_data = json.load(f)
        
        items = items_data.get('items', [])
        print(f"✓ Loaded {len(items)} items", flush=True)
        
        # Check description quality
        items_with_desc = sum(1 for item in items if len(item.get('description', '')) > 50)
        print(f"  Items with descriptions (>50 chars): {items_with_desc}/{len(items)}", flush=True)
        
        # Load company data (optional)
        company_data = {}
        if os.path.exists(company_data_path):
            with open(company_data_path, 'r', encoding='utf-8') as f:
                company_data = json.load(f)
        
        # Find template file
        template_path_options = [
            os.path.join(BASE_DIR, 'offer2_template.docx'),
            os.path.join(BASE_DIR, 'uploads', 'offer2_template.docx'),
        ]
        
        template_path = None
        for path in template_path_options:
            if os.path.exists(path):
                template_path = path
                print(f"✓ Found template: {path}", flush=True)
                break
        
        if not template_path:
            print("✗ Template not found!", flush=True)
            return False
        
        # Copy template to output
        print(f"Copying template to: {output_path}", flush=True)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        shutil.copy2(template_path, output_path)
        print("✓ Template copied", flush=True)
        
        # Open the copied template
        print("Opening copied template...", flush=True)
        doc = Document(output_path)
        
        print(f"Template has {len(doc.sections)} section(s)", flush=True)
        print(f"Template has {len(doc.paragraphs)} paragraph(s)", flush=True)
        print(f"Template has {len(doc.tables)} table(s)", flush=True)
        
        # Clear body content (keep header/footer)
        print("Clearing body content...", flush=True)
        
        for para in doc.paragraphs[:]:
            p_element = para._element
            p_element.getparent().remove(p_element)
        
        for table in doc.tables[:]:
            t_element = table._element
            t_element.getparent().remove(t_element)
        
        print("✓ Body content cleared", flush=True)
        
        # Add new content
        print("\nAdding new content...", flush=True)
        
        # Generate metadata
        today = datetime.now()
        quote_number = f"QT-{today.strftime('%Y-%m%d-%H%M')}"
        quote_date = today.strftime("%B %d, %Y")
        valid_until = (today + timedelta(days=30)).strftime("%B %d, %Y")
        
        # Create helper
        template_helper = Offer3Template()
        template_helper.doc = doc
        
        # 1. Document info
        print("  → Adding document info...", flush=True)
        template_helper.add_document_info_table(quote_number, quote_date, valid_until, "[Customer Name]")
        
        # 2. Pricing table
        print("  → Adding pricing table...", flush=True)
        template_helper.add_pricing_table(items)
        
        # 3. Technical descriptions from AI-extracted JSON
        print("  → Adding technical descriptions from extracted data...", flush=True)
        add_technical_descriptions_from_json(doc, items)
        
        # 4. Commercial terms
        print("  → Adding commercial terms...", flush=True)
        template_helper.add_commercial_terms(company_data)
        
        # Save
        print(f"\nSaving document: {output_path}", flush=True)
        doc.save(output_path)
        
        file_size = os.path.getsize(output_path)
        print(f"✓ Document saved: {file_size:,} bytes", flush=True)
        
        # Summary
        print("\n" + "=" * 60, flush=True)
        print("OFFER 3 GENERATION SUMMARY", flush=True)
        print("=" * 60, flush=True)
        print(f"Total Items: {len(items)}", flush=True)
        print(f"Items with descriptions: {items_with_desc}", flush=True)
        print(f"Quote Number: {quote_number}", flush=True)
        print(f"Valid Until: {valid_until}", flush=True)
        print(f"Source: AI semantic extraction", flush=True)
        print("=" * 60, flush=True)
        
        print("\n✓ OFFER 3 GENERATION COMPLETED SUCCESSFULLY", flush=True)
        
        return True
        
    except Exception as e:
        print(f"\n✗ FATAL ERROR: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False


def add_technical_descriptions_from_json(doc, items):
    """
    Add technical descriptions section using AI-extracted data
    
    For each item with description/specifications:
    - Item heading (item number + name)
    - Description paragraph
    - Specifications (if available)
    - Details (if available)
    """
    
    # Section heading
    heading = doc.add_paragraph()
    run = heading.add_run("Technical Specifications")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    item_counter = 1
    items_added = 0
    
    for item in items:
        description = item.get('description', '').strip()
        specifications = item.get('specifications', '').strip()
        details = item.get('details', '').strip()
        
        # Skip items with no technical content
        if not description and not specifications and not details:
            item_counter += 1
            continue
        
        # Item heading (e.g., "1. DISCONTINUOUS DISTILLATION UNIT C27")
        item_heading = doc.add_paragraph()
        item_name = item.get('item_name', 'Item')
        run = item_heading.add_run(f"{item_counter}. {item_name}")
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Description (main technical content)
        if description:
            desc_para = doc.add_paragraph()
            run = desc_para.add_run(description)
            run.font.size = Pt(11)
            doc.add_paragraph()  # Spacing
        
        # Specifications (structured technical data)
        if specifications:
            spec_heading = doc.add_paragraph()
            run = spec_heading.add_run("Key Specifications:")
            run.font.bold = True
            run.font.size = Pt(11)
            
            spec_para = doc.add_paragraph()
            run = spec_para.add_run(specifications)
            run.font.size = Pt(10)
            doc.add_paragraph()  # Spacing
        
        # Additional details (notes, custom tariff, etc.)
        if details:
            details_para = doc.add_paragraph()
            run = details_para.add_run(details)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
            doc.add_paragraph()  # Spacing
        
        item_counter += 1
        items_added += 1
    
    print(f"  ✓ Added technical descriptions for {items_added} items", flush=True)
    
    if items_added == 0:
        print(f"  ⚠ WARNING: No items had technical descriptions!", flush=True)
        print(f"  ⚠ Check extraction quality in items_offer1.json", flush=True)


if __name__ == "__main__":
    print("Offer 3 Generation Script Started", flush=True)
    
    company_data_path = os.path.join(OUTPUT_FOLDER, "company_data.json")
    items_data_path = os.path.join(OUTPUT_FOLDER, "items_offer1.json")
    output_path = os.path.join(OUTPUT_FOLDER, "final_offer3.docx")
    
    success = generate_offer3(company_data_path, items_data_path, output_path)
    
    if not success:
        print("✗ Generation failed", flush=True)
        sys.exit(1)
    
    print("✓ COMPLETED SUCCESSFULLY", flush=True)
    sys.exit(0)
"""
Pure Python document to PDF conversion (no LibreOffice needed)
Works on any server including Render.com free tier
Converts to PDF to preserve structure for GPT-4 Vision
"""

from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import openpyxl
from PIL import Image
import os
import shutil


def convert_docx_to_pdf_python(docx_path, pdf_path):
    """Convert DOCX to PDF using reportlab (preserves structure)"""
    try:
        print("  Converting DOCX to PDF using Python libraries...", flush=True)
        
        # Read DOCX
        doc = Document(docx_path)
        
        # Create PDF with A4 page size (more standard for business docs)
        pdf = SimpleDocTemplate(
            pdf_path, 
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12,
            spaceBefore=12,
            alignment=TA_LEFT
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#333333'),
            spaceAfter=10,
            spaceBefore=10,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#000000'),
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            leading=14
        )
        
        # Extract content from DOCX
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Escape special XML characters for reportlab
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Detect heading level
                if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Title':
                    p = Paragraph(text, title_style)
                elif paragraph.style.name.startswith('Heading'):
                    p = Paragraph(text, heading_style)
                else:
                    p = Paragraph(text, normal_style)
                
                story.append(p)
                story.append(Spacer(1, 0.1*inch))
        
        # Extract tables with proper styling
        for table_idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text and escape special characters
                    cell_text = cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    row_data.append(cell_text)
                data.append(row_data)
            
            if data:
                # Calculate column widths
                num_cols = len(data[0]) if data else 1
                available_width = 6.5 * inch  # A4 width minus margins
                col_width = available_width / num_cols
                
                # Create table with calculated widths
                t = Table(data, colWidths=[col_width] * num_cols)
                
                # Apply professional styling
                table_style = TableStyle([
                    # Header row styling
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 0), (-1, 0), 12),
                    
                    # Data rows styling
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    
                    # Alternating row colors
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
                    
                    # Grid
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BOX', (0, 0), (-1, -1), 1, colors.black),
                    
                    # Alignment
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ])
                
                t.setStyle(table_style)
                story.append(Spacer(1, 0.2*inch))
                story.append(t)
                story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        pdf.build(story)
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"  ✓ DOCX converted to PDF ({file_size:,} bytes)", flush=True)
            return True
        else:
            print("  ✗ PDF file not created", flush=True)
            return False
        
    except Exception as e:
        print(f"  ✗ Python DOCX conversion failed: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False


def convert_xlsx_to_pdf_python(xlsx_path, pdf_path):
    """Convert XLSX to PDF using openpyxl + reportlab (preserves table structure)"""
    try:
        print("  Converting XLSX to PDF using Python libraries...", flush=True)
        
        # Read Excel
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        
        # Create PDF
        pdf = SimpleDocTemplate(
            pdf_path, 
            pagesize=A4,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Process all sheets
        for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
            sheet = workbook[sheet_name]
            
            # Add sheet name as heading
            if len(workbook.sheetnames) > 1:
                sheet_heading = Paragraph(f"<b>{sheet_name}</b>", styles['Heading1'])
                story.append(sheet_heading)
                story.append(Spacer(1, 0.2*inch))
            
            # Extract data
            data = []
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                # Skip completely empty rows
                if any(cell.strip() for cell in row_data):
                    data.append(row_data)
            
            if data:
                # Calculate column widths based on content
                num_cols = len(data[0]) if data else 1
                available_width = 7.5 * inch
                col_width = available_width / num_cols
                
                # Create table
                t = Table(data, colWidths=[col_width] * num_cols)
                
                # Apply Excel-like styling
                table_style = TableStyle([
                    # Header row
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#217346')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, 0), 10),
                    
                    # Data rows
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('TOPPADDING', (0, 1), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    
                    # Grid
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BOX', (0, 0), (-1, -1), 1, colors.black),
                    
                    # Alternating rows
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E7E6E6')]),
                    
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ])
                
                t.setStyle(table_style)
                story.append(t)
                
                # Add page break between sheets (except for last sheet)
                if sheet_idx < len(workbook.sheetnames) - 1:
                    story.append(PageBreak())
        
        # Build PDF
        pdf.build(story)
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"  ✓ XLSX converted to PDF ({file_size:,} bytes)", flush=True)
            return True
        else:
            print("  ✗ PDF file not created", flush=True)
            return False
        
    except Exception as e:
        print(f"  ✗ Python XLSX conversion failed: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False


def convert_image_to_pdf_python(image_path, pdf_path):
    """Convert image to PDF (wraps image in PDF container)"""
    try:
        print("  Converting image to PDF using Python libraries...", flush=True)
        
        # Open image to get dimensions
        img = Image.open(image_path)
        img_width, img_height = img.size
        
        # Calculate aspect ratio
        aspect = img_height / float(img_width)
        
        # Use A4 size and scale image to fit
        page_width, page_height = A4
        
        # Calculate image size to fit page with margins
        margin = 0.5 * inch
        max_width = page_width - 2 * margin
        max_height = page_height - 2 * margin
        
        if aspect > 1:  # Portrait image
            new_width = min(max_width, img_width)
            new_height = new_width * aspect
            if new_height > max_height:
                new_height = max_height
                new_width = new_height / aspect
        else:  # Landscape image
            new_height = min(max_height, img_height)
            new_width = new_height / aspect
            if new_width > max_width:
                new_width = max_width
                new_height = new_width * aspect
        
        # Create PDF
        pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
        story = []
        
        # Center the image
        img_obj = RLImage(image_path, width=new_width, height=new_height)
        story.append(Spacer(1, margin))
        story.append(img_obj)
        
        pdf.build(story)
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"  ✓ Image converted to PDF ({file_size:,} bytes)", flush=True)
            return True
        else:
            print("  ✗ PDF file not created", flush=True)
            return False
        
    except Exception as e:
        print(f"  ✗ Image conversion failed: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False


def convert_to_pdf_python(input_path, output_path, file_format):
    """
    Main conversion function using pure Python (no LibreOffice)
    Preserves document structure for optimal GPT-4 Vision extraction
    """
    
    print(f"Converting {file_format.upper()} to PDF using Python libraries...", flush=True)
    print(f"  Input: {input_path} ({os.path.getsize(input_path):,} bytes)", flush=True)
    print(f"  Output: {output_path}", flush=True)
    
    try:
        if file_format == 'docx' or file_format == 'doc':
            success = convert_docx_to_pdf_python(input_path, output_path)
            
        elif file_format == 'xlsx' or file_format == 'xls':
            success = convert_xlsx_to_pdf_python(input_path, output_path)
        
        elif file_format in ['png', 'jpg', 'jpeg']:
            success = convert_image_to_pdf_python(input_path, output_path)
        
        elif file_format == 'pdf':
            # Already PDF, just copy
            shutil.copy(input_path, output_path)
            print(f"  ✓ PDF copied (already in correct format)", flush=True)
            success = True
        
        else:
            print(f"  ✗ Unsupported format: {file_format}", flush=True)
            success = False
        
        return success
            
    except Exception as e:
        print(f"  ✗ Conversion error: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False
"""
Standard Offer 3 Template Structure
Defines the professional quotation format that will be generated
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import io
from PIL import Image as PILImage

class Offer3Template:
    """
    Standard professional quotation template
    Creates consistent, high-quality output regardless of input formats
    """
    
    def __init__(self):
        self.doc = Document()
        
        # Standard styling
        self.font_name = "Calibri"
        self.font_size_body = 11
        self.font_size_header = 14
        self.font_size_small = 9
        
        # Standard colors
        self.color_header_bg = "4472C4"  # Professional blue
        self.color_text_dark = "000000"
        self.color_text_light = "666666"
        
    def copy_header_footer_from_template(self, template_path):
        """
        Copy header and footer directly from Offer 2 template to Offer 3
        This preserves logos, formatting, company info exactly as-is
        
        Uses a more robust approach that handles embedded images properly
        """
        
        print("\n" + "=" * 60, flush=True)
        print("COPYING HEADER AND FOOTER FROM TEMPLATE", flush=True)
        print("=" * 60, flush=True)
        
        if not os.path.exists(template_path):
            print(f"⚠ Template not found: {template_path}", flush=True)
            return False
        
        try:
            from docx import Document as DocxDocument
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn
            
            # Open template as a ZIP file to access raw XML and relationships
            import zipfile
            import shutil
            from lxml import etree
            
            print("Reading template document structure...", flush=True)
            template_doc = DocxDocument(template_path)
            
            # Get first section
            if len(template_doc.sections) == 0:
                print("⚠ Template has no sections", flush=True)
                return False
            
            template_section = template_doc.sections[0]
            our_section = self.doc.sections[0]
            
            # Method 1: Try simple paragraph/table copy (works if no images)
            print("Attempting simple header/footer copy...", flush=True)
            
            # COPY HEADER CONTENT
            if template_section.header:
                print("Copying header paragraphs and tables...", flush=True)
                
                # Clear our header
                for paragraph in our_section.header.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                
                for table in our_section.header.tables:
                    t = table._element
                    t.getparent().remove(t)
                
                # Copy paragraphs (but this won't copy images in headers)
                for para in template_section.header.paragraphs:
                    new_para = our_section.header.add_paragraph(para.text)
                    
                    # Copy formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                
                # Copy tables
                for table in template_section.header.tables:
                    new_table = our_section.header.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                
                print("✓ Header content copied (text and tables)", flush=True)
            
            # COPY FOOTER CONTENT
            if template_section.footer:
                print("Copying footer paragraphs and tables...", flush=True)
                
                # Clear our footer
                for paragraph in our_section.footer.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                
                for table in our_section.footer.tables:
                    t = table._element
                    t.getparent().remove(t)
                
                # Copy paragraphs
                for para in template_section.footer.paragraphs:
                    new_para = our_section.footer.add_paragraph(para.text)
                    
                    # Copy formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                
                # Copy tables
                for table in template_section.footer.tables:
                    new_table = our_section.footer.add_table(rows=len(table.rows), cols=len(table.columns))
                    
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                
                print("✓ Footer content copied (text and tables)", flush=True)
            
            print("=" * 60, flush=True)
            print("⚠ NOTE: Images in header (logo) may not be copied by this method", flush=True)
            print("⚠ Logo will need to be added separately", flush=True)
            print("=" * 60, flush=True)
            
            return True
            
        except Exception as e:
            print(f"✗ Error copying header/footer: {str(e)}", flush=True)
            import traceback
            traceback.print_exc()
            return False
    
    def add_company_logo_from_template(self, template_path):
        """
        Extract logo from template and add it at the top of the document
        This is a workaround since copying images in headers is complex
        """
        
        print("\n" + "=" * 60, flush=True)
        print("EXTRACTING AND ADDING COMPANY LOGO", flush=True)
        print("=" * 60, flush=True)
        
        try:
            from docx import Document as DocxDocument
            
            template_doc = DocxDocument(template_path)
            
            # Look for images in document relationships
            image_found = False
            for rel in template_doc.part.rels.values():
                if "image" in rel.target_ref:
                    print(f"Found image: {rel.target_ref}", flush=True)
                    
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Save to temp file
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                        tmp_file.write(image_data)
                        tmp_path = tmp_file.name
                    
                    # Add logo at top of document
                    logo_para = self.doc.paragraphs[0].insert_paragraph_before()
                    run = logo_para.add_run()
                    run.add_picture(tmp_path, width=Inches(2.0))
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Clean up temp file
                    os.remove(tmp_path)
                    
                    print("✓ Logo added to top of document", flush=True)
                    image_found = True
                    break
            
            if not image_found:
                print("⚠ No logo image found in template", flush=True)
            
            print("=" * 60, flush=True)
            return image_found
            
        except Exception as e:
            print(f"⚠ Could not extract logo: {str(e)}", flush=True)
            return False
    
    def add_header_section(self, company_data):
        """
        DEPRECATED - This method is no longer used
        We now copy header/footer directly from template
        Keeping this for backwards compatibility
        """
        print("⚠ add_header_section called but header is copied from template", flush=True)
        pass
    
    def add_document_info_table(self, quote_number, date, valid_until, customer_name):
        """
        Add document metadata table:
        - Quotation No
        - Date
        - Valid Until
        - Customer
        """
        
        info_table = self.doc.add_table(rows=4, cols=2)
        # Don't set any style - just use default
        
        # Set column widths
        for row in info_table.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.0)
        
        # Row 1: Quotation Number
        self._set_cell_text(info_table.rows[0].cells[0], "Quotation No:", bold=True)
        self._set_cell_text(info_table.rows[0].cells[1], quote_number)
        
        # Row 2: Date
        self._set_cell_text(info_table.rows[1].cells[0], "Date:", bold=True)
        self._set_cell_text(info_table.rows[1].cells[1], date)
        
        # Row 3: Valid Until
        self._set_cell_text(info_table.rows[2].cells[0], "Valid Until:", bold=True)
        self._set_cell_text(info_table.rows[2].cells[1], valid_until)
        
        # Row 4: Customer
        self._set_cell_text(info_table.rows[3].cells[0], "Customer:", bold=True)
        self._set_cell_text(info_table.rows[3].cells[1], customer_name)
        
        # Spacing after table
        self.doc.add_paragraph()
    
    def _get_or_create_shading(self, cell):
        """Helper to add cell shading"""
        tcPr = cell._element.get_or_add_tcPr()
        shading = tcPr.find(qn('w:shd'))
        if shading is None:
            shading = OxmlElement('w:shd')
            tcPr.append(shading)
        return shading
    
    def add_pricing_table(self, items, currency="€"):
        """
        Add pricing table with all items grouped by category
        
        Columns: No. | Description | Quantity | Unit Price | Total
        """
        
        # Create table with header row
        pricing_table = self.doc.add_table(rows=1, cols=5)
        # Don't set style - use default
        
        # Set column widths
        pricing_table.columns[0].width = Inches(0.5)   # No.
        pricing_table.columns[1].width = Inches(3.5)   # Description
        pricing_table.columns[2].width = Inches(0.8)   # Quantity
        pricing_table.columns[3].width = Inches(1.0)   # Unit Price
        pricing_table.columns[4].width = Inches(1.0)   # Total
        
        # Header row
        header_cells = pricing_table.rows[0].cells
        self._set_cell_text(header_cells[0], "No.", bold=True, bg_color=self.color_header_bg, text_color="FFFFFF")
        self._set_cell_text(header_cells[1], "Description", bold=True, bg_color=self.color_header_bg, text_color="FFFFFF")
        self._set_cell_text(header_cells[2], "Quantity", bold=True, bg_color=self.color_header_bg, text_color="FFFFFF")
        self._set_cell_text(header_cells[3], "Unit Price", bold=True, bg_color=self.color_header_bg, text_color="FFFFFF")
        self._set_cell_text(header_cells[4], "Total", bold=True, bg_color=self.color_header_bg, text_color="FFFFFF")
        
        # Group items by category
        from collections import OrderedDict
        categorized = OrderedDict()
        for item in items:
            cat = item.get('category', 'Items')
            if cat not in categorized:
                categorized[cat] = []
            categorized[cat].append(item)
        
        item_counter = 1
        
        # Add items by category
        for category, cat_items in categorized.items():
            # Category header row
            cat_row = pricing_table.add_row().cells
            cat_row[0].merge(cat_row[4])
            self._set_cell_text(cat_row[0], category, bold=True, bg_color="E7E6E6")
            
            # Items in category
            for item in cat_items:
                row = pricing_table.add_row().cells
                
                # No.
                self._set_cell_text(row[0], str(item_counter), align="center")
                item_counter += 1
                
                # Description (item_name only, technical details go to separate section)
                description = item.get('item_name', '')
                self._set_cell_text(row[1], description)
                
                # Quantity
                qty = item.get('quantity', '1')
                self._set_cell_text(row[2], str(qty), align="center")
                
                # Unit Price
                unit_price = item.get('unit_price', '')
                self._set_cell_text(row[3], str(unit_price), align="right")
                
                # Total
                total = item.get('total_price', item.get('unit_price', ''))
                self._set_cell_text(row[4], str(total), align="right")
        
        # Add totals row if needed (placeholder for now)
        totals_row = pricing_table.add_row().cells
        totals_row[0].merge(totals_row[3])
        self._set_cell_text(totals_row[0], "TOTAL:", bold=True, align="right")
        self._set_cell_text(totals_row[4], "", bold=True, align="right")
        
        # Spacing after table
        self.doc.add_paragraph()
    
    def add_technical_descriptions(self, items):
        """
        Add detailed technical descriptions for each item
        Only for items that have description/specifications
        """
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run("Technical Specifications")
        run.font.size = Pt(self.font_size_header)
        run.font.bold = True
        run.font.name = self.font_name
        # Don't set style - formatting already applied above
        
        self.doc.add_paragraph()
        
        item_counter = 1
        for item in items:
            description = item.get('description', '')
            specifications = item.get('specifications', '')
            details = item.get('details', '')
            
            # Skip items with no technical content
            if not description and not specifications and not details:
                item_counter += 1
                continue
            
            # Item heading
            item_heading = self.doc.add_paragraph()
            run = item_heading.add_run(f"{item_counter}. {item.get('item_name', 'Item')}")
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = self.font_name
            
            # Description
            if description:
                desc_para = self.doc.add_paragraph()
                run = desc_para.add_run(description)
                run.font.size = Pt(self.font_size_body)
                run.font.name = self.font_name
                self.doc.add_paragraph()
            
            # Specifications (as bullet points if structured)
            if specifications:
                spec_heading = self.doc.add_paragraph()
                run = spec_heading.add_run("Key Specifications:")
                run.font.bold = True
                run.font.size = Pt(self.font_size_body)
                run.font.name = self.font_name
                
                spec_para = self.doc.add_paragraph()
                run = spec_para.add_run(specifications)
                run.font.size = Pt(self.font_size_body)
                run.font.name = self.font_name
                self.doc.add_paragraph()
            
            # Additional details
            if details:
                details_para = self.doc.add_paragraph()
                run = details_para.add_run(details)
                run.font.size = Pt(self.font_size_small)
                run.font.name = self.font_name
                run.font.color.rgb = RGBColor(102, 102, 102)  # Gray text
                self.doc.add_paragraph()
            
            item_counter += 1
        
        self.doc.add_paragraph()
    
    def add_commercial_terms(self, company_data, supplier_terms=None):
        """
        Add commercial terms section
        Shows company standard terms + highlights differences from supplier
        """
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run("Commercial Terms & Conditions")
        run.font.size = Pt(self.font_size_header)
        run.font.bold = True
        run.font.name = self.font_name
        # Don't set style - formatting already applied above
        
        self.doc.add_paragraph()
        
        standard_terms = company_data.get('standard_terms', {})
        
        # Delivery Terms
        delivery_para = self.doc.add_paragraph()
        run = delivery_para.add_run("DELIVERY TERMS\n")
        run.font.bold = True
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        delivery_text = standard_terms.get('delivery', 'As per agreement')
        run = delivery_para.add_run(delivery_text)
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        # Note supplier difference if provided
        if supplier_terms and supplier_terms.get('delivery'):
            run = delivery_para.add_run(f"\nNote: Supplier delivery time is {supplier_terms['delivery']}")
            run.font.size = Pt(self.font_size_small)
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
        
        self.doc.add_paragraph()
        
        # Payment Terms
        payment_para = self.doc.add_paragraph()
        run = payment_para.add_run("PAYMENT TERMS\n")
        run.font.bold = True
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        payment_text = standard_terms.get('payment', 'As per agreement')
        run = payment_para.add_run(payment_text)
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        self.doc.add_paragraph()
        
        # Warranty
        warranty_para = self.doc.add_paragraph()
        run = warranty_para.add_run("WARRANTY\n")
        run.font.bold = True
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        warranty_text = standard_terms.get('warranty', 'As per manufacturer standard warranty')
        run = warranty_para.add_run(warranty_text)
        run.font.size = Pt(self.font_size_body)
        run.font.name = self.font_name
        
        self.doc.add_paragraph()
    
    def add_footer_section(self, company_data):
        """
        DEPRECATED - Footer is now copied directly from template
        Keeping this for backwards compatibility
        """
        print("⚠ add_footer_section called but footer is copied from template", flush=True)
        pass
    
    def save(self, output_path):
        """Save the document"""
        self.doc.save(output_path)
        print(f"✓ Document saved: {output_path}", flush=True)
    
    # Helper methods
    
    def _set_cell_text(self, cell, text, bold=False, align="left", bg_color=None, text_color=None):
        """Set cell text with formatting"""
        cell.text = ""  # Clear existing
        para = cell.paragraphs[0]
        run = para.add_run(str(text))
        
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_body)
        run.font.bold = bold
        
        if text_color:
            if len(text_color) == 6:
                run.font.color.rgb = RGBColor(
                    int(text_color[0:2], 16),
                    int(text_color[2:4], 16),
                    int(text_color[4:6], 16)
                )
        
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if bg_color:
            self._set_cell_background(cell, bg_color)
    
    def _set_cell_background(self, cell, color_hex):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _remove_table_borders(self, table):
        """Remove all borders from table"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
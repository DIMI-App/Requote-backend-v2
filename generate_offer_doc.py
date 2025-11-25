import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import Counter, OrderedDict
import openai

openai.api_key = os.environ.get('OPENAI_API_KEY')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OFFER_2_PATH = os.path.join(BASE_DIR, "offer2_template.docx")
ITEMS_PATH = os.path.join(BASE_DIR, "outputs", "items_offer1.json")
OUTPUT_PATH = os.path.join(BASE_DIR, "outputs", "final_offer1.docx")

def detect_template_language(doc):
    """Detect the language of Offer 2 template"""
    print("=" * 60, flush=True)
    print("DETECTING TEMPLATE LANGUAGE", flush=True)
    print("=" * 60, flush=True)
    
    text_samples = []
    
    for table in doc.tables:
        for row in table.rows[:5]:
            for cell in row.cells:
                text = cell.text.strip()
                if len(text) > 10:
                    text_samples.append(text)
    
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        if len(text) > 10:
            text_samples.append(text)
    
    if not text_samples:
        print("⚠ No text found in template, defaulting to English", flush=True)
        return "EN"
    
    combined_text = " ".join(text_samples[:5])[:500]
    print(f"Text sample: {combined_text[:150]}...", flush=True)
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{
                "role": "user",
                "content": f"""Detect the language of this text. Return ONLY the two-letter ISO 639-1 code in UPPERCASE (e.g., 'EN', 'UK', 'ES', 'DE').

Text: {combined_text}

Language code:"""
            }],
            max_tokens=10,
            temperature=0
        )
        
        detected_lang = response.choices[0].message.content.strip().upper()
        print(f"✓ Detected language: {detected_lang}", flush=True)
        print("=" * 60, flush=True)
        
        return detected_lang if len(detected_lang) == 2 else "EN"
        
    except Exception as e:
        print(f"✗ Language detection error: {str(e)}", flush=True)
        return "EN"

def analyze_document_context(items):
    """Analyze items to understand document context and extract glossary"""
    print("=" * 60, flush=True)
    print("ANALYZING DOCUMENT CONTEXT", flush=True)
    print("=" * 60, flush=True)
    
    all_text = []
    for item in items[:5]:
        all_text.append(item.get('item_name', ''))
        if item.get('details'):
            all_text.append(item.get('details', ''))
    
    context_sample = "\n".join(all_text)
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": f"""Analyze this technical quotation and provide:

1. Industry/Domain (e.g., "Food & Beverage Processing", "Pharmaceutical Equipment")
2. Main Product Category (e.g., "Bottling Line", "Packaging Machinery")
3. Technical Terms Glossary - List 15-20 key technical terms that should NOT be translated
   Include: model numbers, technical terms, brand names, English industry terms
   Format as JSON array: ["term1", "term2", ...]

Quotation sample:
{context_sample[:1000]}

Return ONLY JSON:
{{
  "industry": "...",
  "product_category": "...",
  "technical_glossary": [...]
}}"""
            }],
            max_tokens=600,
            temperature=0
        )
        
        result_json = response.choices[0].message.content.strip()
        if result_json.startswith("```"):
            result_json = result_json.replace("```json", "").replace("```", "").strip()
        
        context = json.loads(result_json)
        
        print(f"✓ Industry: {context.get('industry', 'Unknown')}", flush=True)
        print(f"✓ Product: {context.get('product_category', 'Unknown')}", flush=True)
        print(f"✓ Glossary: {len(context.get('technical_glossary', []))} terms", flush=True)
        if context.get('technical_glossary'):
            print(f"  Terms: {', '.join(context['technical_glossary'][:5])}...", flush=True)
        print("=" * 60, flush=True)
        
        return context
        
    except Exception as e:
        print(f"✗ Context analysis failed: {e}", flush=True)
        return {
            "industry": "Industrial Equipment",
            "product_category": "Machinery",
            "technical_glossary": []
        }

def translate_items_with_context(items, target_lang, context):
    """Translate items using GPT-4o with full context awareness and examples"""
    if target_lang == 'EN' or target_lang == 'EN-US':
        print("Target language is English, no translation needed", flush=True)
        return items
    
    print("=" * 60, flush=True)
    print(f"TRANSLATING TO {target_lang} WITH ENHANCED CONTEXT", flush=True)
    print("=" * 60, flush=True)
    
    lang_map = {
        'UK': 'Ukrainian (Українська)',
        'ES': 'Spanish (Español)',
        'DE': 'German (Deutsch)',
        'FR': 'French (Français)',
        'IT': 'Italian (Italiano)',
        'RU': 'Russian (Русский)',
        'PL': 'Polish (Polski)',
        'PT': 'Portuguese (Português)'
    }
    
    target_language_name = lang_map.get(target_lang, target_lang)
    
    glossary_items = context.get('technical_glossary', [])[:20]
    glossary_text = "\n- ".join(glossary_items) if glossary_items else "None specified"
    
    # Enhanced system prompt with examples
    system_prompt = f"""You are an expert technical translator specializing in industrial equipment quotations for the {context.get('industry', 'manufacturing')} industry.

DOCUMENT CONTEXT:
- Type: Official B2B Technical Quotation
- Industry: {context.get('industry', 'Industrial Equipment')}
- Product: {context.get('product_category', 'Machinery')}
- Target Language: {target_language_name}
- Audience: Professional procurement managers and engineers

MANDATORY TRANSLATION RULES:

1. TECHNICAL TERMS - NEVER TRANSLATE:
   - Model numbers (e.g., "CAN ISO 20/2 S", "VBS MINIDOSE", "TECNA MC24", "RM28")
   - Technical specifications (e.g., "0,33L", "ø15mm", "AISI 304", "AISI 316")
   - Brand names and product codes
   - English technical terms commonly used in industry (e.g., "dummy", "C.I.P.", "kit", "rolls")
   - Measurement units (mm, kg, L, bph, cph, etc.)
   - Acronyms (e.g., "CAN", "LID", "CO₂")
   
   PRESERVE UNTRANSLATED:
   - {glossary_text}

2. CATEGORY NAMES - Use Professional B2B Terminology:
   Examples for Ukrainian:
   - "Main Equipment" → "Основне технологічне обладнання" (NOT "Основне обладнання")
   - "Format Changes" → "Форматні деталі" (NOT "Зміна формату")
   - "Accessories" → "Додаткове обладнання" (NOT "Аксесуари")
   - "Further Options" → "Додаткові опції" (NOT "Подальші варіанти")
   - "Packing" → "Упаковка для транспортування" (NOT just "Упаковка")
   - "CAN FILLER SANITATION" → "Система санітарної обробки наповнювача банок"
   
   Examples for Spanish:
   - "Main Equipment" → "Equipo principal"
   - "Format Changes" → "Cambios de formato"
   - "Accessories" → "Equipamiento adicional"

3. ITEM DESCRIPTIONS:
   - Keep technical specifications in English within translated text
   - Preserve parenthetical technical details exactly
   - Maintain formal business register
   - Use industry-standard terminology
   - Keep model numbers, part names in original language

4. SPECIAL TERMS:
   - "On request" → "На запит" (Ukrainian) / "Bajo petición" (Spanish)
   - "Included" → "Включено" (Ukrainian) / "Incluido" (Spanish)
   - "dummy" → keep as "dummy" (technical term)
   - "kit" → keep as "kit" (industry standard)
   - "rolls" → keep as "rolls" (technical component)

5. QUALITY STANDARDS:
   - Professional, formal business language
   - Consistent terminology throughout entire document
   - Natural phrasing for native B2B readers
   - Technical accuracy over literal translation
   - Preserve all numbers, measurements, and units exactly

EXAMPLES OF CORRECT TRANSLATION TO UKRAINIAN:

Input: "Main Equipment"
Output: "Основне технологічне обладнання"

Input: "CAN FILLER SANITATION\\nSeries of manual closed dummy CANS + washing cam. The cleansing liquid flows throughout the gas evacuation pipes."
Output: "Санітарна обробка наповнювача банок\\nКомплект муляжей закритих банок + промивальний кулачок. Рідина для очищення протікає через трубки відводу газів."

Input: "Equipment for another diameter of can (screw, stars and guides) with SAME LID"
Output: "Форматні деталі для іншого діаметра банки (шнек, зірки та направляючі) з ТАКОЮ Ж КРИШКОЮ"

Input: "Touch-screen panel, colour, multifunction"
Output: "Сенсорна панель, кольорова, багатофункціональна"

Input: "Set of CO₂ regulators in stainless steel, sanitizable, Teflon tube covered in inox, pipes"
Output: "Набір регуляторів CO₂ з нержавіючої сталі, підходить для санітизації, тефлонова трубка з покриттям з нержавіючої сталі, труби"

CRITICAL: Maintain exact JSON structure in your response. Translate category, item_name, and details fields only. Keep all other fields unchanged."""

    try:
        translated_items = []
        batch_size = 6  # Smaller batches for better quality
        
        for i in range(0, len(items), batch_size):
            batch = items[i:i+batch_size]
            
            print(f"  Translating batch {i//batch_size + 1}/{(len(items)-1)//batch_size + 1}...", flush=True)
            
            user_prompt = f"""Translate these {len(batch)} quotation items to {target_language_name}.

REMINDER: This is a {context.get('product_category', 'industrial machinery')} quotation for {context.get('industry', 'professional use')}.

Apply all translation rules from your system instructions:
- Preserve technical terms, model numbers, and specifications
- Use professional B2B terminology for categories
- Keep English technical terms that are industry standard
- Maintain formal register throughout

Input JSON:
{json.dumps(batch, ensure_ascii=False, indent=2)}

Output (translated JSON with same structure):"""
            
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=4500,
                temperature=0.15  # Lower for more consistent terminology
            )
            
            batch_json = response.choices[0].message.content.strip()
            if batch_json.startswith("```"):
                batch_json = batch_json.replace("```json", "").replace("```", "").strip()
            
            batch_translated = json.loads(batch_json)
            translated_items.extend(batch_translated)
            
            print(f"    ✓ Batch {i//batch_size + 1} completed", flush=True)
        
        if len(translated_items) > 0:
            sample_orig = items[0]
            sample_trans = translated_items[0]
            print("\n📊 Translation Sample:", flush=True)
            print(f"  Original category: '{sample_orig.get('category', '')}'", flush=True)
            print(f"  Translated category: '{sample_trans.get('category', '')}'", flush=True)
            print(f"  Original item: '{sample_orig.get('item_name', '')[:70]}'", flush=True)
            print(f"  Translated item: '{sample_trans.get('item_name', '')[:70]}'", flush=True)
        
        print(f"✅ ENHANCED TRANSLATION COMPLETED: {len(translated_items)} items", flush=True)
        print("=" * 60, flush=True)
        
        return translated_items
        
    except Exception as e:
        print(f"✗ Translation failed: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return items

def translate_items(items, target_lang):
    """Main translation entry point with context analysis"""
    if target_lang == 'EN' or target_lang == 'EN-US':
        print("Target language is English, no translation needed", flush=True)
        return items
    
    context = analyze_document_context(items)
    return translate_items_with_context(items, target_lang, context)

def translate_technical_sections(sections, target_lang, context):
    """Translate technical sections with context"""
    if target_lang == 'EN' or target_lang == 'EN-US':
        return sections
    
    print(f"Translating {len(sections)} technical sections to {target_lang}...", flush=True)
    
    lang_map = {
        'UK': 'Ukrainian (Українська)',
        'RU': 'Russian (Русский)',
        'ES': 'Spanish (Español)',
        'DE': 'German (Deutsch)',
        'FR': 'French (Français)',
        'IT': 'Italian (Italiano)'
    }
    
    target_language_name = lang_map.get(target_lang, target_lang)
    glossary_items = context.get('technical_glossary', [])[:20]
    glossary_text = "\n- ".join(glossary_items) if glossary_items else "None specified"
    
    try:
        translated_sections = []
        
        for section in sections:
            prompt = f"""Translate this technical description to {target_language_name}.

CONTEXT:
- Industry: {context.get('industry', 'Industrial Equipment')}
- Product: {context.get('product_category', 'Machinery')}

RULES:
- Preserve technical terms, model numbers, specifications
- DO NOT translate these terms: {glossary_text}
- Maintain professional business language
- Keep numerical values unchanged
- Translate naturally for B2B audience

Original:
Title: {section.get('section_title', '')}
Content: {section.get('content', '')}

Return ONLY JSON:
{{
  "section_title": "translated title",
  "content": "translated content"
}}"""
            
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.2
            )
            
            result_json = response.choices[0].message.content.strip()
            if result_json.startswith("```"):
                result_json = result_json.replace("```json", "").replace("```", "").strip()
            
            translated = json.loads(result_json)
            
            new_section = section.copy()
            new_section['section_title'] = translated.get('section_title', section.get('section_title', ''))
            new_section['content'] = translated.get('content', section.get('content', ''))
            translated_sections.append(new_section)
        
        print(f"✓ Translated {len(translated_sections)} sections", flush=True)
        return translated_sections
        
    except Exception as e:
        print(f"✗ Section translation failed: {str(e)}", flush=True)
        return sections

def get_cell_background_color(cell):
    """Extract background color from cell"""
    try:
        tcPr = cell._element.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill != 'auto':
                return fill
    except:
        pass
    return None

def get_text_color(run):
    """Extract text color from run"""
    try:
        if run.font.color.rgb:
            return str(run.font.color.rgb)
    except:
        pass
    return None

def get_font_name(run):
    """Extract font name from run"""
    try:
        if run.font.name:
            return run.font.name
    except:
        pass
    return None

def get_font_size(run):
    """Extract font size from run"""
    try:
        if run.font.size:
            return run.font.size.pt
    except:
        pass
    return None

def analyze_template_style(doc):
    """Analyze template to detect colors and fonts"""
    style_info = {
        'header_bg_color': None,
        'header_text_color': None,
        'body_text_color': None,
        'primary_font': None,
        'header_font_size': 11,
        'body_font_size': 10
    }
    
    print("=" * 60, flush=True)
    print("ANALYZING TEMPLATE STYLE", flush=True)
    print("=" * 60, flush=True)
    
    bg_colors = []
    text_colors = []
    fonts = []
    font_sizes = []
    
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                bg = get_cell_background_color(cell)
                if bg:
                    bg_colors.append(bg)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        color = get_text_color(run)
                        if color:
                            text_colors.append(color)
                        
                        font = get_font_name(run)
                        if font:
                            fonts.append(font)
                        
                        size = get_font_size(run)
                        if size:
                            font_sizes.append(size)
    
    if bg_colors:
        bg_counter = Counter(bg_colors)
        style_info['header_bg_color'] = bg_counter.most_common(1)[0][0]
        print(f"✓ Header background: #{style_info['header_bg_color']}", flush=True)
    
    if text_colors:
        text_counter = Counter(text_colors)
        most_common = text_counter.most_common(2)
        for color, count in most_common:
            if color.upper() in ['FFFFFF', 'FFFFFFFF']:
                style_info['header_text_color'] = color
                print(f"✓ Header text: #{color}", flush=True)
            else:
                style_info['body_text_color'] = color
                print(f"✓ Body text: #{color}", flush=True)
    
    if fonts:
        font_counter = Counter(fonts)
        style_info['primary_font'] = font_counter.most_common(1)[0][0]
        print(f"✓ Primary font: {style_info['primary_font']}", flush=True)
    
    if font_sizes:
        size_counter = Counter(font_sizes)
        common_sizes = size_counter.most_common(2)
        if len(common_sizes) >= 2:
            sizes_sorted = sorted([s[0] for s in common_sizes], reverse=True)
            style_info['header_font_size'] = sizes_sorted[0]
            style_info['body_font_size'] = sizes_sorted[1]
        elif len(common_sizes) == 1:
            style_info['body_font_size'] = common_sizes[0][0]
        
        print(f"✓ Header font size: {style_info['header_font_size']}pt", flush=True)
        print(f"✓ Body font size: {style_info['body_font_size']}pt", flush=True)
    
    print("=" * 60, flush=True)
    
    return style_info

def set_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def apply_text_style(cell, text, is_header, style_info):
    """Apply font and color styling to cell"""
    cell.text = text
    
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if style_info['primary_font']:
                run.font.name = style_info['primary_font']
            
            if is_header:
                run.font.size = Pt(style_info['header_font_size'])
                run.bold = True
                if style_info['header_text_color']:
                    color_hex = style_info['header_text_color'].replace('#', '')
                    if len(color_hex) == 6:
                        run.font.color.rgb = RGBColor(
                            int(color_hex[0:2], 16),
                            int(color_hex[2:4], 16),
                            int(color_hex[4:6], 16)
                        )
            else:
                run.font.size = Pt(style_info['body_font_size'])
                if style_info['body_text_color']:
                    color_hex = style_info['body_text_color'].replace('#', '')
                    if len(color_hex) == 6:
                        run.font.color.rgb = RGBColor(
                            int(color_hex[0:2], 16),
                            int(color_hex[2:4], 16),
                            int(color_hex[4:6], 16)
                        )

def detect_number_format(table):
    """Detect number format from existing template data"""
    format_info = {
        'thousands_sep': ' ',
        'decimal_sep': '',
        'decimals': 0,
        'currency_symbol': False
    }
    
    print("Analyzing template number format...", flush=True)
    
    for row_idx, row in enumerate(table.rows[:5]):
        for cell in row.cells:
            text = cell.text.strip()
            if re.search(r'\d[\s.,]\d', text):
                print(f"  Sample: '{text}'", flush=True)
                
                if re.search(r'\d+\s\d+$', text):
                    format_info['thousands_sep'] = ' '
                    format_info['decimals'] = 0
                elif re.search(r'\d+\.\d{3},\d{2}', text):
                    format_info['thousands_sep'] = '.'
                    format_info['decimal_sep'] = ','
                    format_info['decimals'] = 2
                elif re.search(r'\d+,\d{3}\.\d{2}', text):
                    format_info['thousands_sep'] = ','
                    format_info['decimal_sep'] = '.'
                    format_info['decimals'] = 2
                
                break
        if format_info['thousands_sep']:
            break
    
    print(f"✓ Detected format: thousands='{format_info['thousands_sep']}', decimals={format_info['decimals']}", flush=True)
    return format_info

def format_price(price_str, format_info):
    """Format price according to template format"""
    if not price_str:
        return ""
    
    price_lower = str(price_str).lower().strip()
    
    if 'included' in price_lower or 'включено' in price_lower or 'incluido' in price_lower:
        return "Included"
    if any(x in price_lower for x in ['on request', 'to be quoted', 'can be offered', 'please inquire', 'на запит', 'за запитом']):
        return "On request"
    
    numeric = re.sub(r'[^\d.,]', '', str(price_str))
    if not numeric:
        return price_str
    
    try:
        clean = numeric.replace('.', '').replace(',', '.')
        value = float(clean)
        
        if format_info['decimals'] == 0:
            formatted = f"{int(value):,}".replace(',', format_info['thousands_sep'])
        else:
            int_part = int(value)
            dec_part = int((value - int_part) * 100)
            int_formatted = f"{int_part:,}".replace(',', format_info['thousands_sep'])
            formatted = f"{int_formatted}{format_info['decimal_sep']}{dec_part:02d}"
        
        return formatted
    except:
        return price_str

def insert_technical_sections(doc, technical_sections, target_language, template_style, context):
    """Insert technical description sections into document"""
    if not technical_sections or len(technical_sections) == 0:
        print("No technical sections to insert", flush=True)
        return
    
    print("=" * 60, flush=True)
    print(f"INSERTING {len(technical_sections)} TECHNICAL SECTIONS", flush=True)
    print("=" * 60, flush=True)
    
    # Find where to insert: before the first table
    first_table_index = None
    for idx, element in enumerate(doc.element.body):
        if element.tag.endswith('tbl'):
            first_table_index = idx
            break
    
    if first_table_index is None:
        print("No table found, inserting at end", flush=True)
    else:
        print(f"Found first table at position {first_table_index}", flush=True)
    
    # Translate technical sections if needed
    if target_language != 'EN':
        print(f"Translating technical sections to {target_language}...", flush=True)
        technical_sections = translate_technical_sections(technical_sections, target_language, context)
    
    # Insert sections before the pricing table
    sections_inserted = 0
    for section in technical_sections:
        location = section.get('page_location', '')
        
        # Only insert sections that should appear before price table
        if location == 'before_price_table' or location == '':
            content_type = section.get('content_type', 'paragraph')
            title = section.get('section_title', '')
            content = section.get('content', '')
            
            if not content:
                continue
            
            # Add section title as heading
            if title:
                heading = doc.add_paragraph(title)
                heading.style = 'Heading 2'
                for run in heading.runs:
                    run.font.bold = True
                    run.font.size = Pt(template_style.get('header_font_size', 12))
                    if template_style.get('primary_font'):
                        run.font.name = template_style['primary_font']
            
            # Add content based on type
            if content_type == 'bullet_list':
                # Split by newlines and add as bullet points
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph(line.strip(), style='List Bullet')
                        for run in p.runs:
                            run.font.size = Pt(template_style.get('body_font_size', 10))
                            if template_style.get('primary_font'):
                                run.font.name = template_style['primary_font']
            else:
                # Add as normal paragraph(s)
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        for run in p.runs:
                            run.font.size = Pt(template_style.get('body_font_size', 10))
                            if template_style.get('primary_font'):
                                run.font.name = template_style['primary_font']
            
            # Add spacing after section
            doc.add_paragraph()
            sections_inserted += 1
            print(f"  ✓ Inserted: {title or 'Untitled section'}", flush=True)
    
    print(f"✓ Inserted {sections_inserted} technical sections", flush=True)
    print("=" * 60, flush=True)

# MAIN EXECUTION
print("=" * 60, flush=True)
print("GENERATE OFFER - SV12 Enhanced Extraction with Technical Sections", flush=True)
print("=" * 60, flush=True)

# Load items AND technical sections
try:
    print(f"Loading data from: {ITEMS_PATH}", flush=True)
    with open(ITEMS_PATH, "r", encoding="utf-8") as f:
        full_data = json.load(f)
    
    items = full_data.get("items", [])
    technical_sections = full_data.get("technical_sections", [])
    
    print(f"✓ Loaded {len(items)} items", flush=True)
    print(f"✓ Loaded {len(technical_sections)} technical sections", flush=True)
    
    # Log SV12 enhanced data if available
    if "extraction_version" in full_data:
        print(f"✓ Extraction version: {full_data.get('extraction_version')}", flush=True)
    if "quality_metrics" in full_data:
        metrics = full_data["quality_metrics"]
        print(f"✓ Quality metrics:", flush=True)
        print(f"  - Description coverage: {metrics.get('description_coverage_percent', 0)}%", flush=True)
        print(f"  - Specifications coverage: {metrics.get('specifications_coverage_percent', 0)}%", flush=True)
        print(f"  - Images detected: {metrics.get('image_detection_percent', 0)}%", flush=True)
    
    if len(items) == 0:
        print("✗ No items found", flush=True)
        exit(1)

except Exception as e:
    print(f"✗ Error loading data: {str(e)}", flush=True)
    exit(1)

# Load template
try:
    print(f"Loading template from: {OFFER_2_PATH}", flush=True)
    doc = Document(OFFER_2_PATH)
    print(f"✓ Template loaded: {len(doc.tables)} tables", flush=True)
except Exception as e:
    print(f"✗ Error loading template: {str(e)}", flush=True)
    exit(1)

if len(doc.tables) == 0:
    print("✗ No tables in template", flush=True)
    exit(1)

# Detect template language
target_language = detect_template_language(doc)

# Analyze document context (for translation glossary)
context = analyze_document_context(items)

# Analyze template style
template_style = analyze_template_style(doc)

# Insert technical sections BEFORE pricing table (if they exist)
if technical_sections and len(technical_sections) > 0:
    insert_technical_sections(doc, technical_sections, target_language, template_style, context)

# Translate items with enhanced context analysis
items = translate_items(items, target_language)

# Find pricing table
best_table = None
max_cols = 0

for idx, table in enumerate(doc.tables):
    cols = len(table.columns)
    print(f"Table {idx + 1}: {len(table.rows)} rows, {cols} columns", flush=True)
    if cols > max_cols and len(table.rows) > 1:
        max_cols = cols
        best_table = table

if best_table is None:
    print("✗ Could not find pricing table", flush=True)
    exit(1)

print(f"✓ Selected table with {len(best_table.columns)} columns", flush=True)

# Detect number format
number_format = detect_number_format(best_table)

# Clear existing data rows
print(f"Clearing {len(best_table.rows) - 1} existing rows...", flush=True)
while len(best_table.rows) > 1:
    best_table._tbl.remove(best_table.rows[1]._tr)

# Group items by category
categorized_items = OrderedDict()
for item in items:
    cat = item.get("category", "Main Items")
    if cat not in categorized_items:
        categorized_items[cat] = []
    categorized_items[cat].append(item)

print(f"✓ Grouped into {len(categorized_items)} categories", flush=True)

# Insert items
item_counter = 1
for category, cat_items in categorized_items.items():
    print(f"  Processing category: {category} ({len(cat_items)} items)", flush=True)
    
    category_row = best_table.add_row().cells
    
    for cell in category_row:
        cell.text = ""
    
    if len(category_row) >= 2:
        apply_text_style(category_row[1], category, True, template_style)
        
        if template_style['header_bg_color']:
            set_cell_background(category_row[1], template_style['header_bg_color'])
    
    for item in cat_items:
        row = best_table.add_row().cells
        
        try:
            if len(row) >= 1:
                apply_text_style(row[0], f"{item_counter}.", False, template_style)
                item_counter += 1
            
            if len(row) >= 2:
                # BUILD COMPLETE DESCRIPTION FROM ALL FIELDS (SV12 Enhancement)
                desc_parts = []
                
                # 1. Item name (always included)
                desc_parts.append(item.get("item_name", ""))
                
                # 2. Full technical description (NEW - main enhancement)
                if item.get("description"):
                    desc_parts.append(item.get("description"))
                
                # 3. Technical specifications (NEW)
                if item.get("specifications"):
                    desc_parts.append(f"Specifications: {item.get('specifications')}")
                
                # 4. Image description (NEW)
                if item.get("has_image") and item.get("image_description"):
                    desc_parts.append(f"[Image: {item.get('image_description')}]")
                
                # 5. Additional details (legacy field, keep for compatibility)
                if item.get("details"):
                    desc_parts.append(item.get("details"))
                
                # Combine all parts with double newline separator
                full_description = "\n\n".join(desc_parts)
                
                apply_text_style(row[1], full_description, False, template_style)
            
            if len(row) >= 3:
                price = format_price(item.get("unit_price", ""), number_format)
                apply_text_style(row[2], price, False, template_style)
            
            if len(row) >= 4:
                apply_text_style(row[3], str(item.get("quantity", "1")), False, template_style)
            
            if len(row) >= 5:
                total = format_price(item.get("total_price", ""), number_format)
                apply_text_style(row[4], total, False, template_style)
            
        except Exception as e:
            print(f"  ✗ Error on item: {str(e)}", flush=True)
            continue

print(f"✓ Inserted all items with enhanced descriptions (SV12)", flush=True)

# Save document
try:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH)
    print(f"✓ Document saved: {OUTPUT_PATH}", flush=True)
    print(f"  File size: {file_size:,} bytes", flush=True)
except Exception as e:
    print(f"✗ Error saving: {str(e)}", flush=True)
    exit(1)

print("=" * 60, flush=True)
print("GENERATION COMPLETE", flush=True)
print("=" * 60, flush=True)
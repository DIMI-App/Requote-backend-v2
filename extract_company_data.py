import os
import sys
import json
import openai
import base64
from docx import Document
from PIL import Image
import io

openai.api_key = os.environ.get('OPENAI_API_KEY')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def extract_logo_from_docx(docx_path):
    """Extract logo image from DOCX header/body"""
    try:
        doc = Document(docx_path)
        
        # Check for images in document
        images = []
        image_index = 0
        
        # Look in document relationships for images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_index += 1
                image_data = rel.target_part.blob
                
                # Convert to base64
                img_base64 = base64.b64encode(image_data).decode('utf-8')
                
                # Detect format
                if image_data.startswith(b'\xff\xd8\xff'):
                    img_format = 'jpeg'
                elif image_data.startswith(b'\x89PNG'):
                    img_format = 'png'
                else:
                    img_format = 'unknown'
                
                images.append({
                    'index': image_index,
                    'data': img_base64,
                    'format': img_format,
                    'size': len(image_data)
                })
        
        if images:
            # Return first image (usually the logo)
            print(f"✓ Found {len(images)} image(s) in document", flush=True)
            return images[0]
        else:
            print("⚠ No logo image found in document", flush=True)
            return None
            
    except Exception as e:
        print(f"✗ Logo extraction failed: {str(e)}", flush=True)
        return None

if __name__ == "__main__":
    print("Company Data Extraction Script Started", flush=True)
    
    offer2_path = os.path.join(BASE_DIR, "offer2_template.docx")
    output_path = os.path.join(OUTPUT_FOLDER, "company_data.json")
    
    if not os.path.exists(offer2_path):
        print(f"✗ Template not found at {offer2_path}", flush=True)
        sys.exit(1)
    
    success = extract_company_data_from_offer2(offer2_path, output_path)
    
    if not success:
        print("✗ Extraction failed", flush=True)
        sys.exit(1)
    
    print("✓ COMPLETED SUCCESSFULLY", flush=True)
    sys.exit(0)
    """Extract company branding and information from Offer 2 template"""
    
    try:
        print("=" * 60, flush=True)
        print("EXTRACTING COMPANY DATA FROM OFFER 2", flush=True)
        print("=" * 60, flush=True)
        
        if not openai.api_key:
            print("✗ OPENAI_API_KEY not set", flush=True)
            return False
        
        print(f"Reading template: {offer2_path}", flush=True)
        
        if not os.path.exists(offer2_path):
            print("✗ Template file not found", flush=True)
            return False
        
        # Extract logo image FIRST
        logo_data = extract_logo_from_docx(offer2_path)
        
        # Read DOCX content directly (more reliable than Vision API)
        print("Reading DOCX content directly...", flush=True)
        doc = Document(offer2_path)
        
        text_content = []
        
        # Extract from paragraphs (first 50 paragraphs should contain company info)
        print("Extracting text from paragraphs...", flush=True)
        for para in doc.paragraphs[:50]:
            if para.text.strip():
                text_content.append(para.text.strip())
        
        # Extract from tables (company info often in header tables)
        print("Extracting text from tables...", flush=True)
        for table_idx, table in enumerate(doc.tables[:10]):
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_content.append(row_text)
        
        combined_text = "\n".join(text_content)
        
        print(f"Extracted {len(combined_text)} characters of text", flush=True)
        print(f"Sample text: {combined_text[:200]}...", flush=True)
        
        # Use GPT-4o for extraction with enhanced prompt
        print("Calling GPT-4o for company data extraction...", flush=True)
        
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": f"""Extract company information from this quotation template.

CRITICAL: Extract information about the COMPANY who owns this template (the seller/quotation creator), NOT about products or customers mentioned in the template.

Look for these patterns:
- Company name: Usually at the top, in headers, or in footer
- Address: Street, city, postal code, country
- Contact: Phone numbers (with country code like +XX), email addresses, website URLs
- Legal info: VAT number, Tax ID, Registration number
- Bank details: IBAN, SWIFT/BIC code, bank name, account holder
- Commercial terms: Delivery timeframes (e.g., "14 working weeks"), payment terms (e.g., "30% advance"), warranty terms

Template content:
{combined_text[:4000]}

Return ONLY valid JSON (no markdown, no backticks):
{{
  "company_name": "Full legal company name",
  "address": "Complete address with city and country",
  "phone": "Phone number with country code",
  "email": "Email address",
  "website": "Website URL",
  "tax_id": "VAT or Tax ID number",
  "registration_no": "Company registration number",
  "bank_details": {{
    "bank_name": "Name of the bank",
    "iban": "IBAN number",
    "swift": "SWIFT/BIC code",
    "account_holder": "Account holder name"
  }},
  "standard_terms": {{
    "delivery": "Delivery timeframe or terms",
    "payment": "Payment terms and conditions",
    "warranty": "Warranty terms"
  }},
  "legal_info": "Any additional legal information, registration details, or certifications"
}}

If a field is not found in the template, use empty string "".
"""
            }],
            max_tokens=2000,
            temperature=0
        )
        
        extracted_json = response.choices[0].message.content.strip()
        
        print("Received response from GPT-4o", flush=True)

        
        # Clean JSON formatting
        if extracted_json.startswith("```json"):
            extracted_json = extracted_json.replace("```json", "").replace("```", "").strip()
        elif extracted_json.startswith("```"):
            extracted_json = extracted_json.replace("```", "").strip()
        
        print("Parsing extracted data...", flush=True)
        company_data = json.loads(extracted_json)
        
        # Add logo data if extracted
        if logo_data:
            company_data['logo'] = {
                'format': logo_data['format'],
                'data': logo_data['data'],
                'size': logo_data['size']
            }
            print(f"✓ Logo included ({logo_data['size']} bytes, {logo_data['format']})", flush=True)
        else:
            company_data['logo'] = None
            print("⚠ No logo found", flush=True)
        
        # Validation and display
        print("\n" + "=" * 60, flush=True)
        print("EXTRACTED COMPANY DATA", flush=True)
        print("=" * 60, flush=True)
        print(f"Company: {company_data.get('company_name', 'N/A')}", flush=True)
        print(f"Address: {company_data.get('address', 'N/A')[:80]}...", flush=True)
        print(f"Phone: {company_data.get('phone', 'N/A')}", flush=True)
        print(f"Email: {company_data.get('email', 'N/A')}", flush=True)
        print(f"Website: {company_data.get('website', 'N/A')}", flush=True)
        print(f"Tax ID: {company_data.get('tax_id', 'N/A')}", flush=True)
        
        bank = company_data.get('bank_details', {})
        print(f"Bank: {bank.get('bank_name', 'N/A')}", flush=True)
        print(f"IBAN: {bank.get('iban', 'N/A')}", flush=True)
        print(f"SWIFT: {bank.get('swift', 'N/A')}", flush=True)
        
        terms = company_data.get('standard_terms', {})
        print(f"Delivery: {terms.get('delivery', 'N/A')}", flush=True)
        print(f"Payment: {terms.get('payment', 'N/A')}", flush=True)
        print(f"Warranty: {terms.get('warranty', 'N/A')}", flush=True)
        print("=" * 60, flush=True)
        
        # Validate that we got at least company name
        if not company_data.get('company_name') or company_data.get('company_name') == '':
            print("⚠ WARNING: Company name not extracted. Extraction may have failed.", flush=True)
            print("⚠ Saving partial data anyway...", flush=True)
        
        # Save to JSON
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(company_data, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Saved to {output_path}", flush=True)
        print("=" * 60, flush=True)
        
        return True
        
    except Exception as e:
        print(f"✗ FATAL ERROR: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        
        # Try to save empty structure so process doesn't completely fail
        try:
            empty_data = {
                "company_name": "",
                "address": "",
                "phone": "",
                "email": "",
                "website": "",
                "tax_id": "",
                "registration_no": "",
                "bank_details": {
                    "bank_name": "",
                    "iban": "",
                    "swift": "",
                    "account_holder": ""
                },
                "standard_terms": {
                    "delivery": "",
                    "payment": "",
                    "warranty": ""
                },
                "legal_info": "",
                "logo": None
            }
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(empty_data, f, indent=2, ensure_ascii=False)
            
            print("⚠ Saved empty company data structure to allow process to continue", flush=True)
        except:
            pass
        
        return False

if __name__ == "__main__":
    print("Company Data Extraction Script Started", flush=True)
    
    offer2_path = os.path.join(BASE_DIR, "offer2_template.docx")
    output_path = os.path.join(OUTPUT_FOLDER, "company_data.json")
    
    if not os.path.exists(offer2_path):
        print(f"✗ Template not found at {offer2_path}", flush=True)
        sys.exit(1)
    
    success = extract_company_data_from_offer2(offer2_path, output_path)
    
    if not success:
        print("✗ Extraction failed", flush=True)
        sys.exit(1)
    
    print("✓ COMPLETED SUCCESSFULLY", flush=True)
    sys.exit(0)
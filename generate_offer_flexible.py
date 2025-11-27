import os
import sys
import json
import subprocess
import openai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import OrderedDict

openai.api_key = os.environ.get('OPENAI_API_KEY')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
ITEMS_PATH = os.path.join(OUTPUT_FOLDER, "items_offer1.json")
TEMPLATE_STRUCTURE_PATH = os.path.join(OUTPUT_FOLDER, "template_structure.json")

# NEW PROMPT 3: FUNCTIONAL MERGING INSTRUCTIONS (Complete from uploaded document)
RECOMPOSITION_PROMPT = """PROMPT 3: FUNCTIONAL MERGING INSTRUCTIONS

================================================================================
CONTEXT
================================================================================

You are creating professional quotation (Offer 3) by merging:
- DATA from Offer 1 (supplier quotation with semantic extraction)
- STRUCTURE from Offer 2 (company template with functional map)

CRITICAL PRINCIPLE: Match function to function, not content to content.

Example:
- Offer 1: distillation equipment with steam pressure specs
- Offer 2: bottling equipment template with voltage specs  
- Both have: "technical specifications section"
- Action: Replace voltage specs with steam specs, keeping format

================================================================================
YOUR TASK
================================================================================

Generate DETAILED, ACTIONABLE INSTRUCTIONS for building final document.
For each functional section in Offer 2, specify exactly how to merge Offer 1 data.

================================================================================
OUTPUT FORMAT
================================================================================

Return JSON with this structure:

{{
  "merge_strategy": {{
    "compatibility_analysis": {{
      "offer1_industry": "from Offer 1 context",
      "offer2_industry": "from Offer 2 template analysis",
      "industries_match": true/false,
      "approach": "functional_mapping",
      "note": "All B2B quotations share common functional sections - we map by function"
    }},
    
    "content_availability": {{
      "offer1_has_pricing": true/false,
      "offer1_has_specs": true/false,
      "offer1_has_description": true/false,
      "offer1_has_images": true/false,
      "offer1_has_commercial_terms": true/false
    }}
  }},

  "section_by_section_instructions": [
    
    {{
      "section_number": 1,
      "function": "header_company_info",
      "offer2_location": "from template map",
      "action": "keep_template_structure",
      "specific_updates": [
        "Update date to: [current or Offer 1 date]",
        "Update offer number to: [specify format]",
        "Keep all template branding unchanged"
      ],
      "instruction": "Preserve entire header. Only update dynamic fields: date, offer number.",
      "style_to_maintain": "Template's exact header formatting"
    }},

    {{
      "section_number": 2,
      "function": "main_equipment_title",
      "offer2_location": "from template map",
      "offer1_source": "main_equipment.name",
      "action": "replace_entire_block",
      "instruction": "Delete current title. Insert Offer 1 equipment name. Maintain template's title format (uppercase, bold, multi-line if needed).",
      "style_to_maintain": {{
        "uppercase": true/false,
        "bold": true/false,
        "alignment": "left/center"
      }}
    }},

    {{
      "section_number": 6,
      "function": "pricing_section",
      "offer2_location": "from template map",
      "offer1_source": "pricing_items[]",
      "offer1_items_count": "X",
      "action": "delete_all_items_and_rebuild",
      "instruction": "Delete existing pricing items. Rebuild using Offer 1 items. Maintain template format exactly (POS numbers, price format, columns).",
      "content_mapping": [
        {{
          "offer1_item": "pricing_items[0]",
          "type": "main_equipment / accessory / packing / option",
          "description": "Full description",
          "quantity": 1,
          "unit_price": 96900,
          "total": 96900,
          "currency": "EUR",
          "output_format": "Exact format matching template"
        }}
      ],
      "style_to_maintain": {{
        "pos_format": "Sequential numbers starting from 1",
        "price_format": "€ X.XXX (space, dot separator)",
        "alignment": "left for text, right for numbers"
      }}
    }}
  ],

  "critical_rules": [
    "NEVER delete or modify template header, footer, company branding",
    "ALWAYS maintain template typography: fonts, sizes, uppercase headers",
    "Match FUNCTION to FUNCTION, not content to content",
    "When industries differ, adapt content to template format",
    "Pricing format MUST match template exactly",
    "POS numbering MUST be sequential starting from 1",
    "Section headers MUST remain uppercase if template uses uppercase"
  ],

  "matching_summary": {{
    "successful_matches": [
      "Offer 1 main_equipment → Offer 2 equipment_title",
      "Offer 1 technical_specifications → Offer 2 specs_section",
      "Offer 1 technical_description → Offer 2 description_section",
      "Offer 1 pricing_items → Offer 2 pricing_section",
      "Offer 1 commercial_terms → Offer 2 commercial_terms"
    ]
  }}
}}

================================================================================
INSTRUCTIONS FOR GENERATING OUTPUT
================================================================================

1. ANALYZE COMPATIBILITY
   - Note if industries match or differ
   - Confirm functional mapping approach
   - Check what content is available in Offer 1

2. FOR EACH SECTION:
   - State the function clearly
   - Show template location and format
   - Show Offer 1 source data
   - Specify exact action
   - Provide detailed instruction
   - Give content mapping examples
   - Specify style rules to maintain

3. BE SPECIFIC:
   - "Replace X with Y" is too vague
   - "Delete lines 1-5, insert Offer 1 item descriptions in format: POS X - DESC - QTY - PRICE, maintaining uppercase and € X.XXX format" is specific

4. PROVIDE EXAMPLES:
   - Show exact output format
   - Show how to combine multiple items
   - Show how to handle edge cases

5. DOCUMENT STYLE:
   - Font names and sizes
   - Uppercase/lowercase rules  
   - Alignment rules
   - Spacing rules
   - Price format rules

================================================================================
REMEMBER
================================================================================

These instructions will be executed by Python code using python-docx library.
Instructions must be:
- ✅ Actionable (can be coded)
- ✅ Specific (exact format specified)
- ✅ Complete (covers all sections)
- ✅ Unambiguous (one clear interpretation)

The goal: Professional quotation that looks like it came from template company,
but contains supplier's equipment and pricing.

================================================================================
DATA TO ANALYZE
================================================================================

OFFER 1 EXTRACTION:
{extraction_json}

OFFER 2 TEMPLATE STRUCTURE:
{template_structure_json}

Now create the complete functional merging instructions:
"""

def load_extraction_data():
    """Load Offer 1 extraction"""
    try:
        with open(ITEMS_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"ERROR loading extraction: {e}", flush=True)
        return None

def load_template_structure():
    """Load Offer 2 template analysis"""
    try:
        with open(TEMPLATE_STRUCTURE_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data.get('structure', {})
    except Exception as e:
        print(f"Template structure not found: {e}", flush=True)
        return None

def run_prompt_2_analysis():
    """Run PROMPT 2 to analyze template structure"""
    try:
        print("=" * 60, flush=True)
        print("RUNNING PROMPT 2 (Template Analysis)", flush=True)
        print("=" * 60, flush=True)
        
        analyze_script = os.path.join(BASE_DIR, 'analyze_offer2_template.py')
        
        if not os.path.exists(analyze_script):
            print(f"ERROR: analyze_offer2_template.py not found", flush=True)
            return False
        
        result = subprocess.run(
            ['python', analyze_script],
            capture_output=True,
            text=True,
            cwd=BASE_DIR,
            timeout=120
        )
        
        if result.stdout:
            print(result.stdout, flush=True)
        if result.stderr:
            print(result.stderr, flush=True)
        
        if result.returncode != 0:
            print("ERROR: PROMPT 2 failed", flush=True)
            return False
        
        if os.path.exists(TEMPLATE_STRUCTURE_PATH):
            print("✓ PROMPT 2 completed successfully", flush=True)
            return True
        else:
            print("ERROR: PROMPT 2 did not create template_structure.json", flush=True)
            return False
        
    except subprocess.TimeoutExpired:
        print("ERROR: PROMPT 2 timed out", flush=True)
        return False
    except Exception as e:
        print(f"ERROR running PROMPT 2: {e}", flush=True)
        import traceback
        traceback.print_exc()
        return False

def normalize_extraction_data(extraction_data):
    """
    CRITICAL FIX: Normalize extraction data to handle both OLD and NEW structures
    This ensures compatibility regardless of which extraction method was used
    """
    print("=" * 60, flush=True)
    print("NORMALIZING EXTRACTION DATA", flush=True)
    print("=" * 60, flush=True)
    
    # Check if this is NEW structure (has extraction_method field)
    extraction_method = extraction_data.get('extraction_method', 'UNKNOWN')
    print(f"  Detection: {extraction_method}", flush=True)
    
    # If NEW structure exists (from PROMPT 1), use it
    if extraction_data.get('pricing_items'):
        print(f"  ✓ Found NEW structure: {len(extraction_data['pricing_items'])} pricing_items", flush=True)
        return extraction_data
    
    # If OLD structure exists (from SV12), convert it to NEW structure
    if extraction_data.get('items'):
        print(f"  ⚠ Found OLD structure: {len(extraction_data['items'])} items", flush=True)
        print("  → Converting OLD structure to NEW structure...", flush=True)
        
        old_items = extraction_data.get('items', [])
        
        # Convert old items to new pricing_items format
        pricing_items = []
        for item in old_items:
            # Extract price value
            unit_price_str = item.get('unit_price', '0')
            total_price_str = item.get('total_price', '0')
            
            # Try to extract numeric value
            import re
            unit_price_numeric = 0
            total_price_numeric = 0
            currency = 'EUR'
            
            # Extract currency symbol
            if '€' in str(unit_price_str):
                currency = 'EUR'
            elif '$' in str(unit_price_str):
                currency = 'USD'
            elif '£' in str(unit_price_str):
                currency = 'GBP'
            
            # Extract numeric values
            numbers = re.findall(r'\d+[\.,]?\d*', str(unit_price_str))
            if numbers:
                try:
                    unit_price_numeric = float(numbers[0].replace(',', ''))
                except:
                    unit_price_numeric = 0
            
            numbers = re.findall(r'\d+[\.,]?\d*', str(total_price_str))
            if numbers:
                try:
                    total_price_numeric = float(numbers[0].replace(',', ''))
                except:
                    total_price_numeric = 0
            
            # Determine item type from category
            category = item.get('category', 'Main Equipment').lower()
            if 'main' in category or 'equipment' in category:
                item_type = 'main_equipment'
            elif 'accessor' in category:
                item_type = 'accessory'
            elif 'pack' in category:
                item_type = 'packing'
            elif 'option' in category:
                item_type = 'option'
            else:
                item_type = 'main_equipment'
            
            # Build new format item
            new_item = {
                'type': item_type,
                'description': item.get('item_name', ''),
                'identified_as_main_because': f"Converted from category: {item.get('category', 'unknown')}",
                'quantity': int(item.get('quantity', 1)) if str(item.get('quantity', 1)).isdigit() else 1,
                'unit': 'pcs',
                'unit_price': unit_price_numeric,
                'total': total_price_numeric,
                'currency': currency,
                'full_description': item.get('description', ''),
                'specifications': item.get('specifications', ''),
                'technical_details': item.get('details', ''),
                'has_image': item.get('has_image', False),
                'image_description': item.get('image_description', '')
            }
            
            pricing_items.append(new_item)
        
        # Build normalized structure
        normalized = {
            'extraction_method': 'CONVERTED_FROM_SV12',
            'context_understanding': {
                'industry': 'industrial_equipment',
                'main_product_category': 'machinery',
                'offer_type': 'single_equipment_with_accessories'
            },
            'main_equipment': {
                'name': pricing_items[0]['description'] if pricing_items else 'Equipment',
                'identified_from': 'First item in converted list',
                'reasoning': 'Converted from SV12 extraction'
            },
            'pricing_items': pricing_items,
            'technical_specifications': [],
            'technical_description': {
                'full_text': '',
                'identified_from': 'Not available in SV12',
                'reasoning': 'Legacy extraction format'
            },
            'images': [],
            'commercial_terms': extraction_data.get('commercial_terms', {}),
            'pricing_summary': {},
            'certifications': [],
            'items': old_items  # Keep for backward compatibility
        }
        
        print(f"  ✓ Converted {len(pricing_items)} items to NEW structure", flush=True)
        print("=" * 60, flush=True)
        
        return normalized
    
    # No valid data found
    print("  ✗ ERROR: No valid extraction data found (neither pricing_items nor items)", flush=True)
    print("=" * 60, flush=True)
    return None

def get_recomposition_plan(extraction_data, template_structure):
    """Get intelligent recomposition plan from GPT using NEW PROMPT 3"""
    try:
        print("=" * 60, flush=True)
        print("GETTING MERGING INSTRUCTIONS (NEW PROMPT 3)", flush=True)
        print("=" * 60, flush=True)
        
        # Simplify data to avoid token limits
        simplified_extraction = {
            "context_understanding": extraction_data.get('context_understanding', {}),
            "main_equipment": extraction_data.get('main_equipment', {}),
            "pricing_items_count": len(extraction_data.get('pricing_items', [])),
            "pricing_items_preview": extraction_data.get('pricing_items', [])[:3],
            "technical_specifications_count": len(extraction_data.get('technical_specifications', [])),
            "technical_specifications_preview": extraction_data.get('technical_specifications', [])[:5],
            "has_technical_description": bool(extraction_data.get('technical_description', {}).get('full_text')),
            "has_commercial_terms": bool(extraction_data.get('commercial_terms', {})),
            "has_images": len(extraction_data.get('images', [])) > 0
        }
        
        prompt = RECOMPOSITION_PROMPT.format(
            extraction_json=json.dumps(simplified_extraction, ensure_ascii=False, indent=2),
            template_structure_json=json.dumps(template_structure, ensure_ascii=False, indent=2)
        )
        
        print(f"Prompt length: {len(prompt)} characters", flush=True)
        print("Calling GPT-4o for functional merging plan...", flush=True)
        
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=3000,
            temperature=0.1
        )
        
        plan_json = response.choices[0].message.content.strip()
        
        # Clean JSON
        if plan_json.startswith("```json"):
            plan_json = plan_json.replace("```json", "").replace("```", "").strip()
        elif plan_json.startswith("```"):
            plan_json = plan_json.replace("```", "").strip()
        
        plan_json = plan_json.strip()
        
        if not plan_json.startswith('{'):
            print(f"ERROR: Response doesn't start with '{{'. Starts with: '{plan_json[:50]}'", flush=True)
            return None
        
        try:
            recomposition_plan = json.loads(plan_json)
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}", flush=True)
            print(f"Response:\n{plan_json[:500]}", flush=True)
            return None
        
        if not isinstance(recomposition_plan, dict):
            print(f"ERROR: Response is not a dictionary", flush=True)
            return None
        
        print("✓ Functional merging plan created", flush=True)
        
        sections = recomposition_plan.get('section_by_section_instructions', [])
        print(f"  Sections to process: {len(sections)}", flush=True)
        
        if sections:
            print("  Section functions:", flush=True)
            for section in sections[:5]:
                print(f"    - {section.get('function', 'unknown')}", flush=True)
        
        print("=" * 60, flush=True)
        
        return recomposition_plan
        
    except Exception as e:
        print(f"ERROR getting recomposition plan: {e}", flush=True)
        import traceback
        traceback.print_exc()
        return None

def execute_recomposition_docx(template_path, extraction_data, recomposition_plan, output_path):
    """Execute recomposition using NEW PROMPT 3 approach"""
    try:
        print("=" * 60, flush=True)
        print("EXECUTING FUNCTIONAL DOCUMENT RECOMPOSITION (NEW PROMPT 3)", flush=True)
        print("=" * 60, flush=True)
        
        doc = Document(template_path)
        
        # Get pricing items from NORMALIZED structure
        pricing_items = extraction_data.get('pricing_items', [])
        
        if not pricing_items or len(pricing_items) == 0:
            print("  ✗ ERROR: No pricing_items found after normalization!", flush=True)
            return False
        
        print(f"STEP 1: Processing {len(pricing_items)} pricing items", flush=True)
        
        # STEP 2: IDENTIFY PRICING TABLE
        print("STEP 2: Identifying pricing table...", flush=True)
        
        pricing_table = None
        tables_to_remove = []
        
        # Find table with MOST columns (likely pricing table)
        max_cols = 0
        for table in doc.tables:
            if len(table.columns) > max_cols and len(table.rows) >= 2:
                if pricing_table:
                    tables_to_remove.append(pricing_table)
                pricing_table = table
                max_cols = len(table.columns)
            else:
                tables_to_remove.append(table)
        
        if not pricing_table:
            print("ERROR: No pricing table found", flush=True)
            return False
        
        print(f"  ✓ Found pricing table: {len(pricing_table.rows)} rows, {len(pricing_table.columns)} cols", flush=True)
        
        # Remove extra tables
        for table in tables_to_remove:
            table._element.getparent().remove(table._element)
        print(f"  ✓ Removed {len(tables_to_remove)} extra tables", flush=True)
        
        # STEP 3: CLEAR PRICING TABLE
        print("STEP 3: Clearing pricing table...", flush=True)
        
        while len(pricing_table.rows) > 0:
            pricing_table._tbl.remove(pricing_table.rows[0]._tr)
        
        # Create new header
        header_row = pricing_table.add_row().cells
        if len(header_row) >= 5:
            header_row[0].text = "Pos"
            header_row[1].text = "Description"
            header_row[2].text = "Unit Price"
            header_row[3].text = "Qty"
            header_row[4].text = "Total"
            
            for cell in header_row[:5]:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
        
        print(f"  ✓ Cleared all rows and created new header", flush=True)
        
        # STEP 4: REMOVE OFFER 2 PARAGRAPHS (except header)
        print("STEP 4: Removing Offer 2 content paragraphs...", flush=True)
        
        paragraphs_to_keep = 2
        removed_count = 0
        
        all_elements = list(doc.element.body)
        
        for i in range(len(all_elements) - 1, -1, -1):
            element = all_elements[i]
            if element.tag.endswith('p'):
                para_position = sum(1 for e in all_elements[:i+1] if e.tag.endswith('p'))
                
                if para_position > paragraphs_to_keep:
                    element.getparent().remove(element)
                    removed_count += 1
        
        print(f"  ✓ Removed {removed_count} paragraphs (kept {paragraphs_to_keep} for header)", flush=True)
        
        # STEP 5: INSERT TECHNICAL CONTENT FROM OFFER 1
        print("STEP 5: Inserting Offer 1 technical content...", flush=True)
        
        tech_description = extraction_data.get('technical_description', {})
        tech_specs = extraction_data.get('technical_specifications', [])
        
        if tech_description.get('full_text') or tech_specs:
            pricing_table = doc.tables[0] if doc.tables else None
            if pricing_table:
                table_element = pricing_table._element
                table_parent = table_element.getparent()
                table_index = list(table_parent).index(table_element)
                
                # Insert technical description
                if tech_description.get('full_text'):
                    desc_para = doc.add_paragraph()
                    run = desc_para.add_run("Technical Description")
                    run.font.bold = True
                    run.font.size = Pt(12)
                    table_parent.insert(table_index, desc_para._element)
                    table_index += 1
                    
                    text_para = doc.add_paragraph(tech_description['full_text'])
                    for run in text_para.runs:
                        run.font.size = Pt(10)
                    table_parent.insert(table_index, text_para._element)
                    table_index += 1
                    
                    print(f"    ✓ Inserted technical description", flush=True)
                
                # Insert technical specifications
                if tech_specs and len(tech_specs) > 0:
                    spec_para = doc.add_paragraph()
                    run = spec_para.add_run("Technical Specifications")
                    run.font.bold = True
                    run.font.size = Pt(12)
                    table_parent.insert(table_index, spec_para._element)
                    table_index += 1
                    
                    for spec in tech_specs[:10]:  # Limit to 10 specs
                        param = spec.get('parameter', '')
                        value = spec.get('value', '')
                        unit = spec.get('unit', '')
                        
                        spec_text = f"{param}: {value}"
                        if unit:
                            spec_text += f" {unit}"
                        
                        spec_item = doc.add_paragraph(spec_text, style='List Bullet')
                        for run in spec_item.runs:
                            run.font.size = Pt(10)
                        table_parent.insert(table_index, spec_item._element)
                        table_index += 1
                    
                    print(f"    ✓ Inserted {len(tech_specs[:10])} technical specifications", flush=True)
                
                # Add spacing
                space = doc.add_paragraph()
                table_parent.insert(table_index, space._element)
        
        # STEP 6: FILL PRICING TABLE WITH OFFER 1 ITEMS
        print("STEP 6: Filling pricing table with Offer 1 items...", flush=True)
        
        # Group items by type
        categorized = OrderedDict()
        for item in pricing_items:
            item_type = item.get('type', 'main_equipment')
            item_type_display = item_type.replace('_', ' ').title()
            
            if item_type_display not in categorized:
                categorized[item_type_display] = []
            categorized[item_type_display].append(item)
        
        item_counter = 1
        for category, cat_items in categorized.items():
            # Category header
            cat_row = pricing_table.add_row().cells
            if len(cat_row) >= 2:
                cat_row[1].text = category
                for para in cat_row[1].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            # Items
            for item in cat_items:
                row = pricing_table.add_row().cells
                
                # Position
                if len(row) >= 1:
                    row[0].text = str(item_counter)
                
                # Description - BUILD FROM MULTIPLE FIELDS
                if len(row) >= 2:
                    desc_parts = []
                    
                    # Main description
                    if item.get('description'):
                        desc_parts.append(item['description'])
                    
                    # Full description (if different from main)
                    if item.get('full_description') and item.get('full_description') != item.get('description'):
                        desc_parts.append(item['full_description'])
                    
                    # Specifications
                    if item.get('specifications'):
                        desc_parts.append(f"Specifications: {item['specifications']}")
                    
                    # Technical details
                    if item.get('technical_details'):
                        desc_parts.append(item['technical_details'])
                    
                    # Image info
                    if item.get('has_image') and item.get('image_description'):
                        desc_parts.append(f"[Image: {item['image_description']}]")
                    
                    full_desc = "\n\n".join(desc_parts)
                    row[1].text = full_desc
                    for para in row[1].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Unit Price
                if len(row) >= 3:
                    unit_price = item.get('unit_price', 0)
                    currency = item.get('currency', 'EUR')
                    if isinstance(unit_price, (int, float)) and unit_price > 0:
                        row[2].text = f"{currency} {unit_price:,.0f}"
                    else:
                        row[2].text = str(unit_price)
                
                # Quantity
                if len(row) >= 4:
                    row[3].text = str(item.get('quantity', 1))
                
                # Total
                if len(row) >= 5:
                    total = item.get('total', 0)
                    currency = item.get('currency', 'EUR')
                    if isinstance(total, (int, float)) and total > 0:
                        row[4].text = f"{currency} {total:,.0f}"
                    else:
                        row[4].text = str(total)
                
                item_counter += 1
        
        print(f"  ✓ Added {item_counter - 1} items to pricing table", flush=True)
        
        # STEP 7: SAVE
        doc.save(output_path)
        file_size = os.path.getsize(output_path)
        
        print("=" * 60, flush=True)
        print(f"✅ DOCUMENT SAVED: {output_path}", flush=True)
        print(f"   File size: {file_size:,} bytes", flush=True)
        print(f"   Items: {item_counter - 1}", flush=True)
        print(f"   Tables: {len(doc.tables)}", flush=True)
        print("=" * 60, flush=True)
        
        return True
        
    except Exception as e:
        print(f"ERROR in execute_recomposition_docx: {e}", flush=True)
        import traceback
        traceback.print_exc()
        return False

def generate_offer_flexible():
    """Main generation function with NEW PROMPT 3"""
    try:
        print("=== FLEXIBLE GENERATION (NEW PROMPT 3) ===", flush=True)
        
        if not openai.api_key:
            print("ERROR: OPENAI_API_KEY not set", flush=True)
            return False
        
        # STEP 1: Load extraction data
        extraction_data = load_extraction_data()
        if not extraction_data:
            print("ERROR: No extraction data", flush=True)
            return False
        
        print(f"✓ Loaded extraction data", flush=True)
        
        # CRITICAL FIX: Normalize data structure
        extraction_data = normalize_extraction_data(extraction_data)
        if not extraction_data:
            print("ERROR: Data normalization failed", flush=True)
            return False
        
        # STEP 2: Load template structure - AUTO-RUN PROMPT 2 IF MISSING
        template_structure = load_template_structure()
        
        if not template_structure:
            print("⚠ Template structure missing, running PROMPT 2 now...", flush=True)
            
            template_docx = os.path.join(BASE_DIR, "offer2_template.docx")
            template_xlsx = os.path.join(BASE_DIR, "offer2_template.xlsx")
            
            if not os.path.exists(template_docx) and not os.path.exists(template_xlsx):
                print("ERROR: No template found", flush=True)
                return False
            
            success = run_prompt_2_analysis()
            
            if not success:
                print("ERROR: PROMPT 2 failed", flush=True)
                return False
            
            template_structure = load_template_structure()
            
            if not template_structure:
                print("ERROR: PROMPT 2 completed but template_structure.json is invalid", flush=True)
                return False
        
        print("✓ Template structure loaded", flush=True)
        
        # STEP 3: Get recomposition plan from GPT using NEW PROMPT 3
        recomposition_plan = get_recomposition_plan(extraction_data, template_structure)
        
        if not recomposition_plan:
            print("⚠ GPT plan failed, using simple fallback", flush=True)
            recomposition_plan = {
                "merge_strategy": {"approach": "functional_mapping"},
                "section_by_section_instructions": [],
                "critical_rules": []
            }
        
        # Save plan for debugging
        plan_path = os.path.join(OUTPUT_FOLDER, "recomposition_plan.json")
        with open(plan_path, 'w', encoding='utf-8') as f:
            json.dump(recomposition_plan, f, indent=2, ensure_ascii=False)
        print(f"✓ Saved plan: {plan_path}", flush=True)
        
        # STEP 4: Find template file
        template_docx = os.path.join(BASE_DIR, "offer2_template.docx")
        template_xlsx = os.path.join(BASE_DIR, "offer2_template.xlsx")
        
        if os.path.exists(template_docx):
            template_path = template_docx
        elif os.path.exists(template_xlsx):
            print("ERROR: XLSX templates not supported in functional system", flush=True)
            return False
        else:
            print("ERROR: No template file found", flush=True)
            return False
        
        # STEP 5: Execute recomposition
        output_path = os.path.join(OUTPUT_FOLDER, "final_offer1.docx")
        success = execute_recomposition_docx(template_path, extraction_data, recomposition_plan, output_path)
        
        if success:
            print("=== FUNCTIONAL RECOMPOSITION COMPLETED (NEW PROMPT 3) ===", flush=True)
            return True
        else:
            print("ERROR: Recomposition failed", flush=True)
            return False
        
    except Exception as e:
        print(f"FATAL ERROR in generate_offer_flexible: {e}", flush=True)
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Functional Generation Script (NEW PROMPT 3)", flush=True)
    
    success = generate_offer_flexible()
    
    if not success:
        print("Generation failed", flush=True)
        sys.exit(1)
    
    print("COMPLETED SUCCESSFULLY", flush=True)
    sys.exit(0)
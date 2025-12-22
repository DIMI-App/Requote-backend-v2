"""
NEW APPROACH: Copy technical content directly from Offer 1
Don't extract → JSON → recreate
Just copy the paragraphs, tables, and images AS-IS
"""

import os
from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy

def copy_technical_content_from_offer1(offer1_path, target_doc, start_after_keyword="TECHNICAL SPECIFICATIONS"):
    """
    Copy ALL content after a keyword from Offer 1 directly to target document
    Preserves tables, images, formatting, everything
    
    Args:
        offer1_path: Path to source Offer 1 document
        target_doc: Target Document object to append to
        start_after_keyword: Start copying after this text is found
    """
    
    print(f"\nCopying technical content from {offer1_path}")
    print(f"Looking for content after: '{start_after_keyword}'")
    
    try:
        source_doc = Document(offer1_path)
        
        # Find where technical content starts
        start_copying = False
        start_index = None
        
        for i, para in enumerate(source_doc.paragraphs):
            if start_after_keyword.lower() in para.text.lower():
                start_copying = True
                start_index = i
                print(f"✓ Found keyword at paragraph {i}: '{para.text[:50]}'")
                break
        
        if not start_copying:
            print(f"⚠ Keyword '{start_after_keyword}' not found, copying all content")
            start_index = 0
        
        # Get all body elements (paragraphs AND tables in correct order)
        body_elements = source_doc.element.body
        
        # Track which element index we're at
        para_index = 0
        table_index = 0
        elements_copied = 0
        
        for element in body_elements:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                if para_index >= start_index:
                    # Copy paragraph
                    new_para = target_doc.add_paragraph()
                    new_para._element = deepcopy(element)
                    elements_copied += 1
                para_index += 1
            
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                if start_index is not None and para_index >= start_index:
                    # Copy table
                    new_table_element = deepcopy(element)
                    target_doc.element.body.append(new_table_element)
                    elements_copied += 1
                table_index += 1
        
        print(f"✓ Copied {elements_copied} elements (paragraphs + tables)")
        return True
        
    except Exception as e:
        print(f"✗ Error copying technical content: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def copy_section_from_document(source_path, target_doc, start_keyword, end_keyword=None, skip_first_para=True):
    """
    More flexible: copy content between two keywords
    
    Args:
        source_path: Source document path
        target_doc: Target document
        start_keyword: Start copying after this
        end_keyword: Stop before this (optional)
        skip_first_para: Skip the paragraph containing start_keyword
    """
    
    print(f"\nCopying section: '{start_keyword}' to '{end_keyword or 'END'}'")
    
    try:
        source_doc = Document(source_path)
        
        # Find start and end positions
        start_idx = None
        end_idx = None
        
        for i, para in enumerate(source_doc.paragraphs):
            text = para.text.strip().lower()
            
            if start_keyword.lower() in text and start_idx is None:
                start_idx = i + (1 if skip_first_para else 0)
                print(f"✓ Start at paragraph {start_idx}")
            
            if end_keyword and end_keyword.lower() in text and start_idx is not None:
                end_idx = i
                print(f"✓ End at paragraph {end_idx}")
                break
        
        if start_idx is None:
            print(f"⚠ Start keyword not found")
            return False
        
        if end_idx is None:
            end_idx = len(source_doc.paragraphs)
            print(f"✓ Copying until end ({end_idx})")
        
        # Copy elements in range
        body_elements = source_doc.element.body
        para_counter = 0
        elements_copied = 0
        
        for element in body_elements:
            if element.tag.endswith('p'):
                if start_idx <= para_counter < end_idx:
                    new_para = target_doc.add_paragraph()
                    new_para._element = deepcopy(element)
                    elements_copied += 1
                para_counter += 1
            
            elif element.tag.endswith('tbl'):
                if start_idx <= para_counter < end_idx:
                    new_table_element = deepcopy(element)
                    target_doc.element.body.append(new_table_element)
                    elements_copied += 1
        
        print(f"✓ Copied {elements_copied} elements")
        return True
        
    except Exception as e:
        print(f"✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # Test
    print("Testing technical content copy...")
    
    source = "/mnt/user-data/uploads/PE_241012A0.docx"
    output = "/tmp/test_copy.docx"
    
    target_doc = Document()
    target_doc.add_heading("Test Document", 0)
    target_doc.add_paragraph("This is test content before technical specs.")
    
    copy_technical_content_from_offer1(source, target_doc)
    
    target_doc.save(output)
    print(f"\nSaved test to {output}")
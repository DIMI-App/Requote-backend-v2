import os
import sys
import json
import openai
from docx import Document
import openpyxl

openai.api_key = os.environ.get('OPENAI_API_KEY')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# NEW PROMPT 2: FUNCTIONAL TEMPLATE ANALYSIS
TEMPLATE_ANALYSIS_PROMPT = """PROMPT 2: FUNCTIONAL TEMPLATE ANALYSIS (OFFER 2)

================================================================================
UNDERSTAND TEMPLATE STRUCTURE
================================================================================

Read this template document and answer:

1. BASIC STRUCTURE
   - How many pages?
   - What are the main sections?
   - Is it formal/informal in tone?
   - Text-heavy or table-heavy?

2. VISUAL STYLE
   - Font types and sizes
   - Colors used (if any)
   - Borders, frames, shading
   - Image formatting style

3. ORGANIZATIONAL PATTERN
   - How is information laid out?
   - What gets emphasized (bold, large text)?
   - Are sections clearly separated or flowing?

Provide your understanding in simple language first.

================================================================================
IDENTIFY FUNCTIONAL SECTIONS
================================================================================

Now map each section by its FUNCTION in the quotation, not by its content.

Every B2B quotation has these universal functions - find them in this template:

FUNCTION 1: HEADER/COMPANY INFO - Company branding, offer identification
FUNCTION 2: MAIN EQUIPMENT TITLE - Primary product name (most prominent)
FUNCTION 3: CAPACITY/PERFORMANCE SUMMARY - Key metrics at glance
FUNCTION 4: TECHNICAL SPECIFICATIONS - Detailed parameters with values
FUNCTION 5: TECHNICAL DESCRIPTION - Narrative explanation of functionality
FUNCTION 6: PRICING SECTION - Items with quantities and prices
FUNCTION 7: ADDITIONAL/OPTIONAL ITEMS - Accessories, options
FUNCTION 8: PACKING/TRANSPORT - Packaging costs/info
FUNCTION 9: EXCLUSIONS - What is NOT included
FUNCTION 10: COMMERCIAL TERMS - Payment, delivery, warranty
FUNCTION 11: CERTIFICATIONS - Standards compliance
FUNCTION 12: IMAGES/PHOTOS - Equipment visuals

For EACH function you find, document:
- LOCATION: Where is it in document?
- FORMAT: How is it structured?
- CURRENT CONTENT: What's there now (examples)
- ACTION: What to do with Offer 1 data
- STYLE: Font, size, alignment, formatting

================================================================================
OUTPUT FORMAT
================================================================================

Return complete analysis as JSON:

{
  "template_structure": {
    "total_pages": 2,
    "style": "formal industrial / casual / mixed",
    "formatting": "plaintext / tables / mixed",
    "fonts": "Font names and sizes",
    "visual_style": "minimal / colorful / branded"
  },

  "functional_sections": [
    {
      "function": "header_company_info",
      "location": "page 1, top",
      "format": "plaintext lines",
      "contains": ["company_name", "offer_number", "date"],
      "action": "keep_as_is",
      "update_only": ["date", "offer_number"],
      "style": {
        "font": "default",
        "size": 11,
        "alignment": "left"
      }
    },
    
    {
      "function": "main_equipment_title",
      "location": "exact location",
      "format": "describe format",
      "recognition_clues": "how to identify this section",
      "current_content_example": "first line of current content",
      "action": "replace_entire_block",
      "replace_with": "main_equipment.name from Offer 1",
      "style": {
        "font": "font name",
        "size": 14,
        "bold": true,
        "uppercase": true
      }
    },
    
    {
      "function": "pricing_section",
      "location": "main body",
      "format": "table with X columns",
      "columns": ["Position", "Description", "Qty", "Unit Price", "Total"],
      "current_items_count": X,
      "price_format": "€ X.XXX or other",
      "action": "delete_all_and_rebuild",
      "style": {
        "uses_pos_numbers": true/false,
        "price_alignment": "right"
      }
    }
  ],

  "critical_formatting_rules": {
    "uppercase_section_headers": true/false,
    "price_format": "exact format with example",
    "uses_tables": true/false,
    "uses_pos_numbering": true/false,
    "line_spacing": "single / double",
    "alignment": "primary alignment style"
  }
}

================================================================================
CRITICAL INSTRUCTIONS
================================================================================

✅ DO:
- Focus on FUNCTION not content
- "Pricing section" is "pricing section" regardless of what's being sold
- Describe FORMAT precisely (alignment, fonts, spacing)
- Note how to RECOGNIZE each section
- Document style rules exactly

❌ DON'T:
- Don't make assumptions based on industry
- Don't mention specific products from template
- Don't describe what SHOULD be there
- Don't invent structure that doesn't exist

The goal: Create a MAP showing WHERE each function is, WHAT FORMAT it uses, and HOW to maintain style.

Now analyze this template:
"""

def analyze_docx_template(template_path):
    """Analyze DOCX template structure"""
    try:
        doc = Document(template_path)
        
        analysis_text = f"\nTEMPLATE TYPE: DOCX\n"
        analysis_text += f"TOTAL PARAGRAPHS: {len(doc.paragraphs)}\n"
        analysis_text += f"TOTAL TABLES: {len(doc.tables)}\n\n"
        
        # First 20 paragraphs
        analysis_text += "FIRST 20 PARAGRAPHS:\n"
        for i, para in enumerate(doc.paragraphs[:20]):
            if para.text.strip():
                style = para.style.name if para.style else "Normal"
                analysis_text += f"  Para {i+1} [{style}]: {para.text[:100]}\n"
        
        # All tables
        analysis_text += f"\nTABLES ({len(doc.tables)} total):\n"
        for idx, table in enumerate(doc.tables):
            analysis_text += f"\n  Table {idx + 1}:\n"
            analysis_text += f"    Rows: {len(table.rows)}, Columns: {len(table.columns)}\n"
            
            if table.rows:
                headers = [cell.text[:30] for cell in table.rows[0].cells]
                analysis_text += f"    Headers: {' | '.join(headers)}\n"
            
            if len(table.rows) > 1:
                analysis_text += f"    Sample rows:\n"
                for row_idx in range(1, min(4, len(table.rows))):
                    row_data = [cell.text[:20] for cell in table.rows[row_idx].cells]
                    analysis_text += f"      Row {row_idx}: {' | '.join(row_data)}\n"
        
        # Last 10 paragraphs
        analysis_text += "\nLAST 10 PARAGRAPHS:\n"
        for i, para in enumerate(doc.paragraphs[-10:]):
            if para.text.strip():
                analysis_text += f"  Para {len(doc.paragraphs) - 10 + i + 1}: {para.text[:100]}\n"
        
        return analysis_text
        
    except Exception as e:
        return f"Error analyzing DOCX: {str(e)}"

def analyze_xlsx_template(template_path):
    """Analyze XLSX template structure"""
    try:
        workbook = openpyxl.load_workbook(template_path)
        sheet = workbook.active
        
        analysis_text = f"\nTEMPLATE TYPE: XLSX\n"
        analysis_text += f"SHEET: {sheet.title}\n"
        analysis_text += f"DIMENSIONS: {sheet.max_row} rows x {sheet.max_column} cols\n\n"
        
        analysis_text += "SHEET STRUCTURE:\n"
        for row_idx in range(1, min(30, sheet.max_row + 1)):
            row = sheet[row_idx]
            row_values = [str(cell.value) if cell.value else '' for cell in row]
            row_text = " | ".join([v[:30] for v in row_values if v.strip()])
            if row_text:
                analysis_text += f"  Row {row_idx}: {row_text}\n"
        
        return analysis_text
        
    except Exception as e:
        return f"Error analyzing XLSX: {str(e)}"

def analyze_template_structure(template_path, output_path):
    """Main analysis function with NEW PROMPT 2"""
    try:
        print("=== STARTING FUNCTIONAL TEMPLATE ANALYSIS (NEW PROMPT 2) ===", flush=True)
        
        if not openai.api_key:
            print("ERROR: OPENAI_API_KEY not set", flush=True)
            return False
        
        if not os.path.exists(template_path):
            print(f"ERROR: Template not found: {template_path}", flush=True)
            return False
        
        file_ext = template_path.lower().split('.')[-1]
        print(f"Template type: {file_ext}", flush=True)
        
        # Get structural information
        if file_ext in ['docx', 'doc']:
            structure_info = analyze_docx_template(template_path)
        elif file_ext in ['xlsx', 'xls']:
            structure_info = analyze_xlsx_template(template_path)
        else:
            print(f"Unsupported template type: {file_ext}", flush=True)
            return False
        
        # Build full prompt
        full_prompt = TEMPLATE_ANALYSIS_PROMPT + "\n\n" + structure_info
        
        print("Calling GPT-4o with NEW PROMPT 2...", flush=True)
        
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": full_prompt}],
            max_tokens=4000,
            temperature=0
        )
        
        analysis_json = response.choices[0].message.content.strip()
        
        # Clean JSON
        if analysis_json.startswith("```json"):
            analysis_json = analysis_json.replace("```json", "").replace("```", "").strip()
        elif analysis_json.startswith("```"):
            analysis_json = analysis_json.replace("```", "").strip()
        
        print("Parsing template analysis...", flush=True)
        template_structure = json.loads(analysis_json)
        
        # Save
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        output_data = {
            "analysis_method": "NEW_PROMPT_2_Functional",
            "template_file": template_path,
            "template_type": file_ext,
            "structure": template_structure
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Functional template analysis saved: {output_path}", flush=True)
        
        # Print summary
        if isinstance(template_structure, dict):
            print("\nTemplate Summary:", flush=True)
            ts = template_structure.get('template_structure', {})
            print(f"  Pages: {ts.get('total_pages', 'N/A')}", flush=True)
            print(f"  Style: {ts.get('style', 'N/A')}", flush=True)
            
            sections = template_structure.get('functional_sections', [])
            print(f"  Functional sections found: {len(sections)}", flush=True)
            for section in sections[:5]:
                print(f"    - {section.get('function', 'unknown')}", flush=True)
        
        print("=== FUNCTIONAL ANALYSIS COMPLETED (NEW PROMPT 2) ===", flush=True)
        return True
        
    except Exception as e:
        print(f"ERROR: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Functional Template Analysis (NEW PROMPT 2)", flush=True)
    
    # Find template
    template_docx = os.path.join(BASE_DIR, "offer2_template.docx")
    template_xlsx = os.path.join(BASE_DIR, "offer2_template.xlsx")
    
    if os.path.exists(template_docx):
        template_path = template_docx
    elif os.path.exists(template_xlsx):
        template_path = template_xlsx
    else:
        print("ERROR: No template found", flush=True)
        sys.exit(1)
    
    output_path = os.path.join(OUTPUT_FOLDER, "template_structure.json")
    
    success = analyze_template_structure(template_path, output_path)
    sys.exit(0 if success else 1)